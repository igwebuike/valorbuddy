from __future__ import annotations

import hashlib
import io
import json
import os
import re
import secrets
from datetime import datetime, timedelta, timezone
from pathlib import Path
from urllib.parse import quote_plus
from typing import Any, List, Optional
import asyncio
import logging
import time

import httpx
from google import genai
from google.genai import types
import jwt
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, EmailStr, Field
from sqlalchemy import Boolean, Column, DateTime, ForeignKey, Integer, JSON, String, Text, create_engine, func, inspect, text
from sqlalchemy.orm import Session, declarative_base, relationship, sessionmaker
from app.agentic import AGENT_CATALOG as VOS_AGENT_CATALOG, TOOL_CATALOG as VOS_TOOL_CATALOG, CORE_PRINCIPLE as VOS_CORE_PRINCIPLE, route_goal, build_fallback_plan, build_agent_prompt, agent_opening

APP_NAME = os.getenv("APP_NAME", "ValorBuddy Enterprise API")
ENVIRONMENT = os.getenv("ENVIRONMENT", "production")
DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./valorbuddy.db")
SECRET_KEY = os.getenv("SECRET_KEY", "dev-change-me-valorbuddy")
ALGORITHM = os.getenv("ALGORITHM", "HS256")
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "10080"))
CORS_ORIGINS = os.getenv("CORS_ORIGINS", "*")
# Prefer GOOGLE_API_KEY so the existing Render variable continues to work.
# GEMINI_API_KEY remains a backward-compatible alias.
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
GEMINI_API_KEY = GOOGLE_API_KEY
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.6-flash")
GEMINI_PLANNER_MODEL = os.getenv("GEMINI_PLANNER_MODEL", GEMINI_MODEL)
GOOGLE_CLOUD_PROJECT = os.getenv("GOOGLE_CLOUD_PROJECT", "")
GOOGLE_CLOUD_LOCATION = os.getenv("GOOGLE_CLOUD_LOCATION", "global")
GOOGLE_GENAI_USE_VERTEXAI = os.getenv("GOOGLE_GENAI_USE_VERTEXAI", "false").lower() == "true"
ENABLE_GOOGLE_SEARCH_GROUNDING = os.getenv("ENABLE_GOOGLE_SEARCH_GROUNDING", "true").lower() == "true"
AI_TIMEOUT_SECONDS = int(os.getenv("AI_TIMEOUT_SECONDS", "35"))
ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "eugene.ebem@gmail.com").lower().strip()
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "")
logger = logging.getLogger("valorbuddy.ai")
GOOGLE_MAPS_API_KEY = os.getenv("GOOGLE_PLACES_API_KEY") or os.getenv("GOOGLE_MAPS_API_KEY")
GOOGLE_CALENDAR_ENABLED = os.getenv("GOOGLE_CALENDAR_ENABLED", "false").lower() == "true"
DATA_DIR = Path(os.getenv("DATA_DIR", "/tmp/valorbuddy"))
UPLOAD_DIR = DATA_DIR / "uploads"
DATA_DIR.mkdir(parents=True, exist_ok=True)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}
engine = create_engine(DATABASE_URL, pool_pre_ping=True, connect_args=connect_args)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True)
    email = Column(String(255), unique=True, index=True, nullable=False)
    password_hash = Column(String(255), nullable=False)
    role = Column(String(50), nullable=False, default="veteran")
    is_active = Column(Boolean, nullable=False, default=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    profile = relationship("UserProfile", back_populates="user", uselist=False)


class UserProfile(Base):
    __tablename__ = "user_profiles"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    first_name = Column(String(120), nullable=False)
    last_name = Column(String(120), nullable=True)
    branch = Column(String(80), nullable=False, default="Army")
    city = Column(String(120), nullable=False, default="")
    state = Column(String(80), nullable=False, default="TX")
    rank = Column(String(120), nullable=True, default="")
    service_status = Column(String(80), nullable=False, default="Veteran")
    service_start_year = Column(String(10), nullable=True, default="")
    service_end_year = Column(String(10), nullable=True, default="")
    deployment_history = Column(Text, nullable=True, default="")
    va_rating = Column(String(30), nullable=True, default="")
    accessibility_needs = Column(JSON, nullable=False, default=list)
    preferred_music_genres = Column(JSON, nullable=False, default=list)
    interests = Column(JSON, nullable=False, default=list)
    preferred_tone = Column(String(120), nullable=False, default="calm, practical, encouraging")
    companion_mode = Column(Boolean, nullable=False, default=True)
    profile_data = Column(JSON, nullable=False, default=dict)
    military_mos = Column(String(160), nullable=False, default="")
    military_job_title = Column(String(255), nullable=False, default="")
    military_experience = Column(Text, nullable=False, default="")
    civilian_career_goal = Column(String(255), nullable=True, default="")
    business_interest = Column(Text, nullable=True, default="")
    military_specialty_description = Column(Text, nullable=True, default="")
    years_of_service = Column(Integer, nullable=True)
    security_clearance = Column(String(100), nullable=True, default="")
    highest_education = Column(String(255), nullable=True, default="")
    civilian_certifications = Column(Text, nullable=True, default="")
    linkedin_url = Column(Text, nullable=True, default="")
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    user = relationship("User", back_populates="profile")


class AuthToken(Base):
    __tablename__ = "auth_tokens"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    token = Column(Text, nullable=False)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class Conversation(Base):
    __tablename__ = "conversations"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    source = Column(String(80), nullable=False, default="web")
    title = Column(String(255), nullable=False, default="ValorBuddy conversation")
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class Message(Base):
    __tablename__ = "messages"
    id = Column(Integer, primary_key=True)
    conversation_id = Column(Integer, ForeignKey("conversations.id"), nullable=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    role = Column(String(50), nullable=False)
    content = Column(Text, nullable=False)
    metadata_json = Column(JSON, nullable=False, default=dict)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class Memory(Base):
    __tablename__ = "memories"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    title = Column(String(255), nullable=False)
    note = Column(Text, nullable=True)
    tags = Column(JSON, nullable=False, default=list)
    image_url = Column(String(500), nullable=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class Reminder(Base):
    __tablename__ = "reminders"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    title = Column(String(255), nullable=False)
    date = Column(String(80), nullable=True)
    time = Column(String(80), nullable=True)
    when_text = Column(String(255), nullable=True)
    note = Column(Text, nullable=True)
    status = Column(String(50), nullable=False, default="active")
    calendar_event_id = Column(String(255), nullable=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    filename = Column(String(255), nullable=False)
    doc_type = Column(String(100), nullable=False, default="general")
    file_url = Column(String(500), nullable=True)
    extracted_text = Column(Text, nullable=True)
    ai_summary = Column(Text, nullable=True)
    analysis_json = Column(JSON, nullable=False, default=dict)
    status = Column(String(40), nullable=False, default="processed")
    processed_at = Column(DateTime(timezone=True), nullable=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class ActivitySearch(Base):
    __tablename__ = "activity_searches"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    city = Column(String(120), nullable=False)
    state = Column(String(80), nullable=False)
    query = Column(String(255), nullable=False)
    provider = Column(String(120), nullable=False, default="Google Places")
    live = Column(Boolean, nullable=False, default=False)
    results = Column(JSON, nullable=False, default=list)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class MusicFavorite(Base):
    __tablename__ = "music_favorites"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    title = Column(String(255), nullable=False)
    url = Column(String(500), nullable=True)
    mood = Column(String(80), nullable=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class AdminAuditLog(Base):
    __tablename__ = "admin_audit_logs"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=True, index=True)
    action = Column(String(255), nullable=False)
    details = Column(Text, nullable=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class CareerDocument(Base):
    __tablename__ = "career_documents"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    document_type = Column(String(40), nullable=False, index=True)
    title = Column(String(255), nullable=False)
    target = Column(String(255), nullable=False, default="")
    content = Column(Text, nullable=False)
    source_profile = Column(JSON, nullable=False, default=dict)
    status = Column(String(40), nullable=False, default="draft")
    version = Column(Integer, nullable=False, default=1)
    ai_generated = Column(Boolean, nullable=False, default=True)
    source_document_id = Column(Integer, ForeignKey("documents.id"), nullable=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120_000).hex()
    return f"pbkdf2_sha256${salt}${digest}"


def verify_password(password: str, stored: str) -> bool:
    try:
        algo, salt, digest = stored.split("$", 2)
        if algo != "pbkdf2_sha256":
            return False
        candidate = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 120_000).hex()
        return secrets.compare_digest(candidate, digest)
    except Exception:
        return False


def create_access_token(user: User) -> str:
    exp = datetime.now(timezone.utc) + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    return jwt.encode({"sub": str(user.id), "email": user.email, "role": user.role, "exp": exp}, SECRET_KEY, algorithm=ALGORITHM)


def get_current_user(authorization: str | None = Header(default=None), db: Session = Depends(get_db)) -> User:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Authentication required")
    token = authorization.split(" ", 1)[1].strip()
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload["sub"])
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")
    user = db.get(User, user_id)
    if not user or not user.is_active:
        raise HTTPException(status_code=401, detail="Invalid user")
    return user


def get_optional_user(authorization: str | None = Header(default=None), db: Session = Depends(get_db)) -> Optional[User]:
    try:
        return get_current_user(authorization, db)
    except Exception:
        return None


def admin_required(user: User = Depends(get_current_user)) -> User:
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user



class RegisterRequest(BaseModel):
    email: EmailStr
    password: str = Field(min_length=6)
    first_name: str
    last_name: str = ""
    rank: str = ""
    branch: str = "Army"
    service_status: str = "Veteran"
    service_start_year: str = ""
    service_end_year: str = ""
    deployment_history: str = ""
    va_rating: str = ""
    city: str = ""
    state: str = ""
    interests: List[str] = []
    accessibility_needs: List[str] = []
    preferred_music_genres: List[str] = []


class ProfileUpdate(BaseModel):
    first_name: str
    last_name: str = ""
    rank: str = ""
    branch: str = "Army"
    service_status: str = "Veteran"
    service_start_year: str = ""
    service_end_year: str = ""
    deployment_history: str = ""
    va_rating: str = ""
    city: str = ""
    state: str = ""
    interests: List[str] = []
    accessibility_needs: List[str] = []
    preferred_music_genres: List[str] = []
    profile_data: dict[str, Any] = {}
    military_mos: str = ""
    military_job_title: str = ""
    military_experience: str = ""
    civilian_career_goal: str = ""
    business_interest: str = ""
    military_specialty_description: str = ""
    years_of_service: int | None = None
    security_clearance: str = ""
    highest_education: str = ""
    civilian_certifications: str = ""
    linkedin_url: str = ""


class LoginRequest(BaseModel):
    email: EmailStr
    password: str


class ProfileOut(BaseModel):
    id: int
    email: str
    role: str
    first_name: str
    last_name: str | None = ""
    rank: str = ""
    branch: str
    service_status: str = "Veteran"
    service_start_year: str = ""
    service_end_year: str = ""
    deployment_history: str = ""
    va_rating: str = ""
    city: str
    state: str
    interests: List[str] = []
    accessibility_needs: List[str] = []
    preferred_music_genres: List[str] = []
    profile_data: dict[str, Any] = {}
    military_mos: str = ""
    military_job_title: str = ""
    military_experience: str = ""
    civilian_career_goal: str = ""
    business_interest: str = ""
    military_specialty_description: str = ""
    years_of_service: int | None = None
    security_clearance: str = ""
    highest_education: str = ""
    civilian_certifications: str = ""
    linkedin_url: str = ""


class LoginResponse(BaseModel):
    token: str
    user: ProfileOut


class CompanionRequest(BaseModel):
    message: str
    conversation_id: int | None = None
    mode: str = "companion"
    lat: Optional[float] = None
    lng: Optional[float] = None
    timezone: str = ""


class ReminderIn(BaseModel):
    title: str
    date: str = ""
    time: str = ""
    when_text: str = ""
    note: str = ""


class MemoryIn(BaseModel):
    title: str
    note: str = ""
    tags: List[str] = []
    image_url: str | None = None


class VapiActionRequest(BaseModel):
    intent: str = "general"
    query: str = ""
    message: str = ""
    first_name: str = ""
    email: str = ""
    branch: str = ""
    city: str = ""
    state: str = ""
    lat: Optional[float] = None
    lng: Optional[float] = None
    title: str = ""
    date: str = ""
    time: str = ""
    memory: str = ""
    mood: str = "calm"
    user_type: str = "Veteran"
    context_items: List[dict[str, Any]] = []


class BranchUpdate(BaseModel):
    branch: str


class CareerGenerateIn(BaseModel):
    document_type: str = Field(pattern="^(resume|cover_letter|business_plan|career_plan)$")
    target: str = Field(min_length=2, max_length=255)
    notes: str = Field(default="", max_length=4000)

class MissionFeedbackIn(BaseModel):
    rating: int | None = Field(default=None, ge=1, le=5)
    useful: bool | None = None
    comment: str = Field(default="", max_length=2000)


def profile_out(user: User) -> ProfileOut:
    p = user.profile
    return ProfileOut(
        id=user.id, email=user.email, role=user.role,
        first_name=p.first_name if p else "Veteran", last_name=p.last_name if p else "",
        rank=p.rank if p else "", branch=p.branch if p else "Army",
        service_status=p.service_status if p else "Veteran",
        service_start_year=p.service_start_year if p else "", service_end_year=p.service_end_year if p else "",
        deployment_history=p.deployment_history if p else "", va_rating=p.va_rating if p else "",
        city=p.city if p else "", state=p.state if p else "", interests=p.interests if p else [],
        accessibility_needs=(p.accessibility_needs or []) if p else [], preferred_music_genres=(p.preferred_music_genres or []) if p else [],
        profile_data=p.profile_data if p and p.profile_data else {},
        military_mos=(p.military_mos or (p.profile_data or {}).get("mos","")) if p else "",
        military_job_title=(p.military_job_title or "") if p else "",
        military_experience=(p.military_experience or "") if p else "",
        civilian_career_goal=(p.civilian_career_goal or "") if p else "",
        business_interest=(p.business_interest or "") if p else "",
        military_specialty_description=(p.military_specialty_description or "") if p else "",
        years_of_service=p.years_of_service if p else None,
        security_clearance=(p.security_clearance or "") if p else "",
        highest_education=(p.highest_education or "") if p else "",
        civilian_certifications=(p.civilian_certifications or "") if p else "",
        linkedin_url=(p.linkedin_url or "") if p else ""
    )


def _genai_client():
    """Create a Google Gen AI client for API-key or Vertex AI authentication."""
    if GOOGLE_GENAI_USE_VERTEXAI:
        kwargs = {"vertexai": True, "location": GOOGLE_CLOUD_LOCATION}
        if GOOGLE_CLOUD_PROJECT:
            kwargs["project"] = GOOGLE_CLOUD_PROJECT
        if GEMINI_API_KEY:
            kwargs["api_key"] = GEMINI_API_KEY
        return genai.Client(**kwargs)
    return genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None


SYSTEM_INSTRUCTION = """You are ValorBuddy, a trusted digital battle buddy for the entire military community.

You support veterans, active-duty service members, retired service members, National Guard, Reserve members, transitioning service members, military spouses, children, dependents, caregivers, Gold Star families, surviving spouses, wounded warriors, and military families from every race, ethnicity, nationality, gender, religion, disability status, background, and branch of service.

Treat every person with dignity, respect, empathy, and professionalism. Never make assumptions about combat history, disability status, VA eligibility, finances, beliefs, health conditions, or personal experiences.

MISSION
Help the user accomplish real tasks, save time, reduce stress, understand options, and take the next useful action. Do not merely describe what you could do.

PERSONALITY
Be warm, calm, capable, practical, encouraging, concise, confident, humble, conversational, and military-aware. Never sound like ChatGPT, a search engine, documentation, or a scripted call center. Do not call yourself an AI assistant unless directly asked. Use the user's name naturally when appropriate.

INTENT FIRST
Identify the user's primary intent and stay focused on it. If the user asks about activities or events, discuss only activities or events. Do not mention documents, reminders, benefits, or unrelated features. If the user asks about restaurants, discuss restaurants only. If the user asks about benefits, discuss benefits only.

Ask at most one focused clarification question when it is truly necessary. If sufficient information exists, act immediately. Never end with a generic question such as “What would you like to do?” Make a useful recommendation or ask a specific next-step question.

LIVE LOCATION
For “near me” requests, current GPS coordinates are authoritative. Never default to Dallas or any other city. Never use a saved city when live GPS exists. If GPS is unavailable for a “near me” request, ask: “What city and state are you in so I can search near you?”

LIVE INFORMATION
Never invent events, places, addresses, business hours, weather, traffic, news, prices, jobs, discounts, schedules, or VA office hours. Use live tools when available. Clearly distinguish live results from general guidance.

TOOL BEHAVIOR
Use only the tools needed for the user's exact request. Combine tool results into one natural answer. Never expose internal reasoning, planner output, tool names, backend systems, or implementation details.

CONVERSATION MEMORY
Use recent conversation naturally. Resolve references such as “it,” “that place,” “the first one,” or “is it open?” from prior context. Do not repeat information the user already provided.

LOCAL RESULTS
Give no more than three strong options unless the user asks for more. Summarize why each option fits. End with one specific next step such as directions, distance, hours, registration details, or filtering by today/free/family-friendly.

MENTAL WELLNESS AND SAFETY
You are not a therapist, doctor, psychiatrist, psychologist, crisis counselor, lawyer, or financial advisor. Never diagnose PTSD, depression, anxiety, or another condition. Never provide clinical treatment or promise eligibility or legal outcomes.

If someone appears stressed or overwhelmed, respond calmly, validate without exaggeration, offer one simple grounding step, offer one practical next action, and ask one gentle follow-up question.

If the user expresses intent to harm themselves or someone else, or says they may not be safe, respond immediately: “In the U.S., please call or text 988 and press 1 for the Veterans Crisis Line now. If there is immediate danger, call 911 or go to the nearest emergency department.” Encourage immediate human support and stay calm.

RESPONSE STYLE
Keep most responses under 220 words unless the user asks for detail. Be direct, natural, specific, and action-oriented. Never say “My next useful step would be,” “I can help by,” or list unrelated capabilities. Actually help.

PROFILE INTELLIGENCE
Use the member profile as context, not as decoration. Personalize recommendations using service branch, rank, service status, years served, deployment history, VA disability rating, location, interests, accessibility needs, family role, and saved preferences. Never expose sensitive profile details unnecessarily. Ask for missing profile details only when they materially improve the answer.

PRIORITY SUPPORT AREAS
Provide practical, current guidance for travel safety, veteran-friendly housing and credit preparation, directions and trip ideas, general investment education, vehicle purchasing, veteran-owned small businesses, military discounts, veteran hiring companies, education, healthcare, VA forms, claims, facilities, caregivers, spouses, dependents, and transition support. For financial, legal, medical, housing, credit, and investment topics, give educational guidance and clearly state when a licensed or accredited professional is needed.

VA FORMS AND PROGRAMS
Use official VA information when discussing forms or programs. Identify the likely form, explain its purpose in plain English, link or direct the user to the official source when available, and provide a checklist of information commonly needed. Do not claim to submit restricted VA forms unless an authorized VA API integration is configured and the user has completed required consent and identity verification.

Every response should leave the user better informed, more confident, or one step closer to completing the goal."""


def _extract_text(response: Any) -> str:
    text = getattr(response, "text", None)
    if text:
        return text.strip()
    try:
        parts = response.candidates[0].content.parts
        return "".join(getattr(part, "text", "") or "" for part in parts).strip()
    except Exception:
        return ""


async def gemini_reply(prompt: str, fallback: str, *, grounded: bool = False, json_mode: bool = False, model: str | None = None) -> str:
    """Google Gen AI SDK wrapper with optional Google Search grounding and JSON output."""
    client = _genai_client()
    if not client:
        return fallback
    tools = [types.Tool(google_search=types.GoogleSearch())] if grounded and ENABLE_GOOGLE_SEARCH_GROUNDING else None
    config = types.GenerateContentConfig(
        system_instruction=SYSTEM_INSTRUCTION,
        temperature=0.35 if json_mode else 0.62,
        max_output_tokens=1400,
        tools=tools,
        response_mime_type="application/json" if json_mode else None,
    )
    def call():
        return client.models.generate_content(model=model or GEMINI_MODEL, contents=prompt, config=config)
    try:
        response = await asyncio.wait_for(asyncio.to_thread(call), timeout=AI_TIMEOUT_SECONDS)
        return _extract_text(response) or fallback
    except Exception as exc:
        logger.exception("Gemini request failed: %s", exc)
        return fallback


async def plan_request(message: str, context: dict[str, Any]) -> dict[str, Any]:
    """Ask Gemini to plan tool use. Keyword routing is only a fail-safe fallback."""
    schema = {
        "intent": "general|local_search|live_web|benefits|create_reminder|save_memory|music|briefing|profile|document_question",
        "needs_location": False,
        "needs_places": False,
        "needs_google_search": False,
        "needs_clarification": False,
        "clarification_question": "",
        "default_query": "",
        "search_query": "",
        "response_goal": "",
    }
    prompt = f"""Return only valid JSON matching this shape: {json.dumps(schema)}

Identify ONE primary intent and select the minimum tools required.
- Use local_search with needs_places=true for nearby activities, veteran events, VFW, American Legion, restaurants, VA facilities, stores, parks, directions, or any 'near me' request.
- Use live_web with needs_google_search=true for current weather, traffic, road conditions, news, current policies, changing benefits rules, jobs, discounts, prices, public event schedules, or other time-sensitive facts not reliably covered by Places.
- Use benefits for benefits guidance.
- Use create_reminder only when the user explicitly asks to create/save a reminder.
- Use save_memory only when the user explicitly asks to remember/save something.
- Use music for music requests.

Stay strictly focused on the user's request. An activities request must not expand into documents, reminders, benefits, or unrelated features.
Ask one clarification question only when the request cannot be acted on accurately without it. When clarification is needed, set needs_clarification=true, write one natural clarification_question, and provide a useful default_query that can be executed if the user does not respond. Do not request clarification when the user already asked clearly for nearby veteran activities or events.

Context: {json.dumps(context, default=str)}
User message: {message}"""
    raw = await gemini_reply(prompt, json.dumps(schema), json_mode=True, model=GEMINI_PLANNER_MODEL)
    try:
        plan = json.loads(raw)
        return {**schema, **plan}
    except Exception:
        fallback_intent = infer_intent(message)
        mapping = {"find_local_veteran_activities":"local_search","search_benefits":"benefits","get_today_briefing":"briefing","get_user_profile":"profile","suggest_music":"music"}
        return {**schema, "intent": mapping.get(fallback_intent, fallback_intent), "needs_places": fallback_intent == "find_local_veteran_activities", "needs_location": fallback_intent == "find_local_veteran_activities", "search_query": message}

def clean_text(text: str) -> str:
    return " ".join((text or "").strip().split())


def human_list(items: list[dict[str, Any]], limit: int = 3) -> str:
    parts = []
    for idx, item in enumerate(items[:limit], 1):
        name = item.get("title") or item.get("name") or "Option"
        loc = item.get("location") or item.get("type") or "nearby"
        rating = f" — rating {item.get('rating')}" if item.get("rating") else ""
        parts.append(f"{idx}. {name} ({loc}){rating}")
    return " ".join(parts)


async def reverse_geocode_location(lat: float | None, lng: float | None) -> dict[str, str]:
    """Resolve browser coordinates into city/state when Google key is available."""
    if lat is None or lng is None or not GOOGLE_MAPS_API_KEY:
        return {"city": "", "state": ""}
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get(
                "https://maps.googleapis.com/maps/api/geocode/json",
                params={"latlng": f"{lat},{lng}", "key": GOOGLE_MAPS_API_KEY},
            )
            r.raise_for_status()
            data = r.json()
        city = ""
        state = ""
        for res in data.get("results", []):
            for c in res.get("address_components", []):
                types = c.get("types", [])
                if "locality" in types or "postal_town" in types:
                    city = c.get("long_name", city)
                if "administrative_area_level_1" in types:
                    state = c.get("short_name", state)
            if city and state:
                break
        return {"city": city, "state": state}
    except Exception:
        return {"city": "", "state": ""}


def _fallback_local_cards(city: str, state: str, query: str) -> list[dict[str, Any]]:
    location_label = f"{city}, {state}" if city and state else "your area"
    maps_location = quote_plus(f"{city} {state}".strip() or "near me")
    clean_query = clean_text(query) or "veteran friendly places"
    fallback_map = {
        "coffee": [
            {"title": "Veteran-friendly coffee nearby", "location": f"Near {location_label}", "type": "Coffee", "description": "Fallback card until live Google Places returns results. Use this for the flow, but verify live options before visiting.", "maps_url": f"https://www.google.com/maps/search/veteran+friendly+coffee+near+{maps_location}"},
            {"title": "American Legion coffee social search", "location": location_label, "type": "Veteran community", "description": "Good for a quick meetup, conversation, or networking with other veterans and families.", "maps_url": f"https://www.google.com/maps/search/American+Legion+near+{maps_location}"},
        ],
        "clinic": [
            {"title": "Nearest VA clinic search", "location": f"Near {location_label}", "type": "VA care", "description": "Open map results for VA clinics and resource centers near the veteran or family member.", "maps_url": f"https://www.google.com/maps/search/VA+clinic+near+{maps_location}"},
            {"title": "Vet Center / counseling resource search", "location": location_label, "type": "Support", "description": "For official support, use VA.gov or call the facility before going.", "maps_url": f"https://www.google.com/maps/search/Vet+Center+near+{maps_location}"},
        ],
        "parks": [
            {"title": "Quiet park or walking trail", "location": f"Near {location_label}", "type": "Wellness", "description": "A simple low-pressure reset option for veterans, spouses, kids, dependents, caregivers, or family.", "maps_url": f"https://www.google.com/maps/search/parks+near+{maps_location}"},
            {"title": "Lake, trail, or outdoor space", "location": location_label, "type": "Outdoor", "description": "Good for fresh air, family time, or decompression.", "maps_url": f"https://www.google.com/maps/search/trails+near+{maps_location}"},
        ],
        "food": [
            {"title": "Mission BBQ / veteran-friendly restaurant search", "location": f"Near {location_label}", "type": "Food", "description": "Look for veteran-friendly restaurants and military discount spots.", "maps_url": f"https://www.google.com/maps/search/veteran+discount+restaurants+near+{maps_location}"},
            {"title": "Veteran-owned restaurant search", "location": location_label, "type": "Food", "description": "Support veteran-owned businesses nearby.", "maps_url": f"https://www.google.com/maps/search/veteran+owned+restaurant+near+{maps_location}"},
        ],
        "family": [
            {"title": "Family-friendly veteran activity search", "location": f"Near {location_label}", "type": "Family", "description": "Options suitable for spouses, kids, dependents, caregivers, and family members.", "maps_url": f"https://www.google.com/maps/search/family+friendly+veteran+events+near+{maps_location}"},
            {"title": "Museums, parks, and community events", "location": location_label, "type": "Kids and family", "description": "Low-pressure local outings for the whole family.", "maps_url": f"https://www.google.com/maps/search/museums+parks+family+events+near+{maps_location}"},
        ],
    }
    q = clean_query.lower()
    if any(k in q for k in ["spouse", "dependent", "kid", "kids", "child", "children", "family", "caregiver"]):
        return fallback_map["family"]
    if any(k in q for k in ["coffee", "breakfast", "cafe"]):
        return fallback_map["coffee"]
    if any(k in q for k in ["clinic", "hospital", "doctor", "va"]):
        return fallback_map["clinic"]
    if any(k in q for k in ["park", "walk", "trail", "outdoor"]):
        return fallback_map["parks"]
    if any(k in q for k in ["food", "restaurant", "bbq", "lunch", "dinner"]):
        return fallback_map["food"]
    return [
        {"title": "VFW or American Legion post", "location": f"Near {location_label}", "type": "Veteran community", "description": "Community, networking, and veteran-friendly events.", "maps_url": f"https://www.google.com/maps/search/VFW+American+Legion+near+{maps_location}"},
        {"title": "Veteran and family-friendly meetup search", "location": location_label, "type": "Social", "description": "A simple option for connection without pressure.", "maps_url": f"https://www.google.com/maps/search/veteran+family+events+near+{maps_location}"},
        {"title": "VA resource or benefits office", "location": f"Near {location_label}", "type": "Benefits", "description": "Useful for official VA questions or referrals.", "maps_url": f"https://www.google.com/maps/search/VA+benefits+office+near+{maps_location}"},
        {"title": "Outdoor family reset spot", "location": f"Near {location_label}", "type": "Wellness", "description": "Park, trail, or quiet outdoor option for veterans, spouses, kids, dependents, caregivers, and family.", "maps_url": f"https://www.google.com/maps/search/parks+near+{maps_location}"},
    ]


async def google_places(city: str = "", state: str = "", query: str = "", lat: float | None = None, lng: float | None = None) -> tuple[bool, list[dict[str, Any]], dict[str, Any]]:
    """Search live Google Places using browser lat/lng first, then city/state. Never silently defaults to Dallas."""
    clean_query = clean_text(query) or "veteran friendly events"
    resolved_city = clean_text(city)
    resolved_state = clean_text(state)
    source = "profile" if resolved_city else "missing"

    if lat is not None and lng is not None:
        rg = await reverse_geocode_location(lat, lng)
        resolved_city = rg.get("city") or resolved_city
        resolved_state = rg.get("state") or resolved_state
        source = "browser_location"

    location_meta = {"city": resolved_city, "state": resolved_state, "lat": lat, "lng": lng, "source": source}

    if not GOOGLE_MAPS_API_KEY:
        return False, _fallback_local_cards(resolved_city, resolved_state, clean_query), {**location_meta, "error": "GOOGLE_PLACES_API_KEY not configured"}

    try:
        params = {"query": clean_query, "key": GOOGLE_MAPS_API_KEY}
        if lat is not None and lng is not None:
            params.update({"location": f"{lat},{lng}", "radius": 25000})
        elif resolved_city and resolved_state:
            params["query"] = f"{clean_query} near {resolved_city}, {resolved_state}"
        else:
            return False, [], {**location_meta, "error": "location_required"}
        async with httpx.AsyncClient(timeout=15) as client:
            r = await client.get("https://maps.googleapis.com/maps/api/place/textsearch/json", params=params)
            r.raise_for_status()
            data = r.json()
        status = data.get("status")
        if status not in ("OK", "ZERO_RESULTS"):
            fb = _fallback_local_cards(resolved_city, resolved_state, clean_query)
            return False, [{**x, "description": f"Google Places returned {status}. Check key, billing, Places API, and restrictions."} for x in fb], {**location_meta, "error": status}
        results = []
        seen = set()
        async with httpx.AsyncClient(timeout=15) as detail_client:
            for item in data.get("results", [])[:8]:
                name = item.get("name") or "Local option"
                if name.lower() in seen:
                    continue
                seen.add(name.lower())
                address = item.get("formatted_address", f"{resolved_city}, {resolved_state}".strip(", "))
                place_id = item.get("place_id")
                reviews = []
                phone = ""
                website = ""
                opening_hours = item.get("opening_hours", {})
                if place_id and len(results) < 3:
                    try:
                        detail_response = await detail_client.get(
                            "https://maps.googleapis.com/maps/api/place/details/json",
                            params={
                                "place_id": place_id,
                                "fields": "formatted_phone_number,website,opening_hours,reviews,url",
                                "reviews_sort": "most_relevant",
                                "key": GOOGLE_MAPS_API_KEY,
                            },
                        )
                        detail = detail_response.json().get("result", {})
                        phone = detail.get("formatted_phone_number", "")
                        website = detail.get("website", "")
                        opening_hours = detail.get("opening_hours") or opening_hours
                        reviews = [
                            {
                                "author": review.get("author_name", "Google reviewer"),
                                "rating": review.get("rating"),
                                "text": clean_text(review.get("text", ""))[:240],
                                "time_description": review.get("relative_time_description", ""),
                            }
                            for review in detail.get("reviews", [])[:2]
                            if clean_text(review.get("text", ""))
                        ]
                    except Exception:
                        reviews = []
                rating = item.get("rating")
                review_total = item.get("user_ratings_total")
                type_text = ", ".join(item.get("types", [])[:3]).replace("_", " ")
                explanation = f"A nearby {type_text or 'community option'} that matches your request."
                if rating:
                    explanation += f" It has a {rating} Google rating"
                    if review_total:
                        explanation += f" from {review_total} reviews"
                    explanation += "."
                results.append({
                    "title": name,
                    "location": address,
                    "type": type_text,
                    "rating": rating,
                    "review_count": review_total,
                    "description": explanation,
                    "assistant_explanation": explanation,
                    "reviews": reviews,
                    "phone": phone,
                    "website": website,
                    "open_now": opening_hours.get("open_now") if isinstance(opening_hours, dict) else None,
                    "place_id": place_id,
                    "maps_url": f"https://www.google.com/maps/search/?api=1&query={quote_plus(name + ' ' + address)}" + (f"&query_place_id={place_id}" if place_id else ""),
                })
        if results:
            return True, results, location_meta
        return False, _fallback_local_cards(resolved_city, resolved_state, clean_query), {**location_meta, "error": "ZERO_RESULTS"}
    except Exception as exc:
        fb = _fallback_local_cards(resolved_city, resolved_state, clean_query)
        return False, [{**x, "description": f"Google Places call failed: {type(exc).__name__}. Check Render env vars and API restrictions."} for x in fb], {**location_meta, "error": type(exc).__name__}


def benefits_lookup(query: str, state: str, branch: str) -> dict[str, Any]:
    q = (query or "benefits").lower()
    items = []
    if any(x in q for x in ["spouse", "dependent", "wife", "husband", "child", "children", "survivor", "caregiver", "family"]):
        items.append({"title": "Spouse, dependent, survivor, and caregiver pathways", "summary": "ValorBuddy can help families understand education, survivor, caregiver, healthcare, and benefit-support pathways in plain English. Eligibility depends on service history and VA rules.", "next_step": "Create a family-access profile, gather DD214/benefit letters, then verify official eligibility on VA.gov or with an accredited VSO.", "assistant_explanation": "This pathway helps families understand which programs may apply and which documents to gather first.", "community_note": "Veteran families commonly say the hardest part is knowing where to start; a document checklist and accredited VSO review can reduce confusion."})
    if any(x in q for x in ["education", "school", "gi", "tuition", "chapter", "dea"]):
        items.append({"title": "Education benefits / GI Bill / DEA starting point", "summary": "Review Post-9/11 GI Bill, transfer rules, Chapter 35 DEA for eligible dependents, school certification, and housing allowance basics.", "next_step": "Gather service records, school/program details, and check the official VA education portal.", "assistant_explanation": "Start by matching the education goal to the correct VA education chapter and confirming school certification.", "community_note": "Many veterans recommend confirming benefit months and housing allowance rules before enrolling."})
    if any(x in q for x in ["disability", "claim", "rating", "compensation", "appeal"]):
        items.append({"title": "VA disability compensation and claim support", "summary": "ValorBuddy can organize evidence, questions, appointments, and plain-English checklists. It does not decide eligibility or replace an accredited representative.", "next_step": "Collect medical/service evidence and speak with a VSO or VA-accredited representative.", "assistant_explanation": "A strong starting point is an organized evidence list, current diagnoses from qualified professionals, and a clear timeline of service-connected events.", "community_note": "Veterans often find accredited VSO support useful for checking forms and evidence before submission."})
    if any(x in q for x in ["home", "loan", "mortgage"]):
        items.append({"title": "VA home loan pathway", "summary": "VA-backed home loans may support buying, refinancing, or repairing a home for eligible veterans and some surviving spouses.", "next_step": "Check Certificate of Eligibility and talk with a VA-approved lender.", "assistant_explanation": "Confirm eligibility, estimate an affordable monthly payment, and compare VA-approved lenders before selecting a property.", "community_note": "Veteran homebuyers frequently recommend comparing lender fees and not assuming every lender offers the same VA loan terms."})
    if any(x in q for x in ["health", "clinic", "medical", "mental", "doctor"]):
        items.append({"title": "VA healthcare and local care navigation", "summary": "Find VA clinics, Vet Centers, community care questions, and appointment reminders. For urgent or crisis needs, call emergency services or 988 then press 1.", "next_step": "Use VA.gov or local VA facility contacts for official enrollment and appointment details.", "assistant_explanation": "ValorBuddy can help locate the closest facility and prepare questions, while the VA confirms enrollment and care options.", "community_note": "Veterans often suggest bringing a medication list, records, and written questions to appointments."})
    if not items:
        items = [
            {"title": "Benefits command center", "summary": "Common categories include healthcare, disability compensation, education, home loan, employment, pension, caregiver, survivor, spouse, and dependent benefits.", "next_step": "Choose a category and ValorBuddy will build a plain-English checklist.", "assistant_explanation": "Pick the benefit area that matters most and receive a focused eligibility and document checklist.", "community_note": "Veterans commonly recommend working one benefit category at a time and keeping copies of every submission."},
            {"title": "Family access", "summary": "Spouses and dependents can use ValorBuddy to organize documents, reminders, resources, and benefit questions connected to the veteran's journey.", "next_step": "Create the appropriate family profile and gather key documents.", "assistant_explanation": "Family pathways vary, so identifying the relationship and benefit goal makes the guidance more accurate.", "community_note": "Military families often say a shared checklist helps everyone understand deadlines and missing documents."},
        ]
    return {"disclaimer": "Informational only. Use VA.gov or a VA-accredited representative for official guidance.", "items": items, "state": state, "branch": branch}


def music_suggestions(mood: str, branch: str) -> list[dict[str, str]]:
    mood_l = (mood or "calm").lower()
    if "patriotic" in mood_l or "military" in mood_l:
        return [{"title": "Patriotic instrumental playlist", "url": "https://www.youtube.com/results?search_query=patriotic+instrumental+music", "mood": mood}, {"title": f"{branch} cadence and heritage music", "url": f"https://www.youtube.com/results?search_query={quote_plus(branch)}+military+cadence", "mood": mood}]
    if "gospel" in mood_l:
        return [{"title": "Calming gospel playlist", "url": "https://www.youtube.com/results?search_query=calming+gospel+playlist", "mood": mood}]
    if "country" in mood_l:
        return [{"title": "Classic country calm mix", "url": "https://www.youtube.com/results?search_query=classic+country+calm+playlist", "mood": mood}]
    return [{"title": "Calm instrumental focus", "url": "https://www.youtube.com/results?search_query=calm+instrumental+music", "mood": mood}, {"title": "Relaxing old school classics", "url": "https://www.youtube.com/results?search_query=relaxing+old+school+classics", "mood": mood}]



def event_choice_payload(first_name: str) -> dict[str, Any]:
    choices = [
        {"label": "Veteran social events", "query": "veteran social events near me"},
        {"label": "VFW or American Legion", "query": "VFW American Legion events near me"},
        {"label": "Family-friendly activities", "query": "family friendly veteran activities near me"},
        {"label": "Outdoor activities", "query": "veteran outdoor activities parks trails fishing near me"},
        {"label": "Live music and entertainment", "query": "veteran friendly live music entertainment near me"},
        {"label": "Coffee and breakfast meetups", "query": "veteran coffee breakfast meetup near me"},
        {"label": "Volunteer opportunities", "query": "veteran volunteer opportunities near me"},
        {"label": "Fitness and recreation", "query": "veteran fitness recreation activities near me"},
        {"label": "Museums and military history", "query": "military museums veteran history attractions near me"},
        {"label": "Career and networking", "query": "veteran career networking events near me"},
        {"label": "Support and wellness groups", "query": "veteran peer support wellness groups near me"},
        {"label": "Free events today", "query": "free veteran events today near me"},
    ]
    labels = "; ".join(f"{i+1}. {x['label']}" for i, x in enumerate(choices))
    response = (
        f"Absolutely, {first_name}. Here are some easy choices: {labels}. "
        "Say the number or the type you want. If you are not sure, say ‘pick for me’ and I’ll start with the best-rated options happening closest to you."
    )
    return {"response": response, "intent": "event_choices", "data": {"choices": choices, "awaiting_choice": True}}


def is_event_choice_request(message: str) -> bool:
    t = (message or "").lower()
    event_words = any(x in t for x in ("event", "events", "activity", "activities", "things to do"))
    choice_words = any(x in t for x in ("example", "examples", "choose", "choices", "options", "types", "what kind", "list"))
    return event_words and choice_words

def infer_intent(text: str) -> str:
    t = (text or "").lower()
    if any(x in t for x in ["remind", "reminder", "appointment", "call the va", "schedule", "tomorrow", "next week"]):
        return "create_reminder"
    if any(x in t for x in ["remember", "memory", "save this", "log this", "journal"]):
        return "save_memory"
    if any(x in t for x in ["benefit", "claim", "gi bill", "disability", "home loan", "va loan", "spouse", "dependent", "survivor", "caregiver", "family access"]):
        return "search_benefits"
    if any(x in t for x in ["event", "activity", "vfw", "american legion", "near me", "places", "coffee", "park", "restaurant", "bbq", "food", "clinic", "va facility", "gym", "fishing"]):
        return "find_local_veteran_activities"
    if any(x in t for x in ["music", "song", "playlist", "play something"]):
        return "suggest_music"
    if any(x in t for x in ["briefing", "today", "how is my day", "what should i do", "plan my day"]):
        return "get_today_briefing"
    if any(x in t for x in ["who am i", "my profile", "profile", "branch"]):
        return "get_user_profile"
    if any(x in t for x in ["travel", "trip", "route", "driving", "road safety", "hotel", "housing", "house", "mortgage", "credit", "investment", "invest", "car", "auto", "vehicle", "veteran owned", "veteran-owned", "military discount", "veteran discount", "hire veterans", "veteran jobs", "company hiring", "va form", "forms", "small business"]):
        return "live_web"
    return "general"


async def route_valorbuddy_message(
    *, text: str, first_name: str, branch: str, city: str = "", state: str = "",
    lat: float | None = None, lng: float | None = None, user_type: str = "Veteran",
    user: Optional[User] = None, db: Optional[Session] = None, explicit_intent: str = "general",
    title: str = "", date: str = "", time: str = "", memory: str = "", mood: str = "calm",
    context_items: Optional[list[dict[str, Any]]] = None
) -> dict[str, Any]:
    """Agentic router: decides which tool to call, gathers data, then composes a human answer."""
    message = clean_text(text)
    context_items = context_items or []
    lower_message = message.lower()
    ownership_followup = any(phrase in lower_message for phrase in [
        "veteran owned", "veteran-owned", "are they owned", "why did you choose",
        "why were they chosen", "how are these events", "are these events",
        "what makes them veteran", "how did you match"
    ])
    if ownership_followup and context_items:
        names = [clean_text(x.get("title")) for x in context_items[:3] if clean_text(x.get("title"))]
        named = ", ".join(names)
        return {
            "response": (
                f"I cannot verify that {named or 'those locations'} are veteran-owned from the available Google listing data. "
                "I selected them because their names, categories, or descriptions indicate that they serve veterans or connect people with veteran resources. "
                "They are locations or organizations—not automatically confirmed scheduled events. For an actual event, I should verify a date, organizer, venue, and event page before recommending it. "
                "Choose Today, This weekend, or a specific date and I’ll narrow the search to confirmed event-style results."
            ),
            "intent": "explain_previous_results",
            "data": {
                "items": context_items[:6],
                "choices": [
                    {"label": "Today", "query": "confirmed veteran events today near me"},
                    {"label": "This weekend", "query": "confirmed veteran events this weekend near me"},
                    {"label": "Veteran-owned only", "query": "verified veteran-owned businesses near me"},
                    {"label": "VFW / Legion events", "query": "scheduled VFW American Legion events near me"}
                ],
                "preserve_results": True
            }
        }
    if is_event_choice_request(message):
        return event_choice_payload(first_name)
    # GPS is authoritative. Reverse-geocode once so every planner and response uses the current area.
    if lat is not None and lng is not None:
        current = await reverse_geocode_location(lat, lng)
        city = current.get("city") or city
        state = current.get("state") or state
    recent_messages = []
    recent_memories = []
    recent_reminders = []
    if user and db:
        recent_messages = db.query(Message).filter(Message.user_id == user.id).order_by(Message.id.desc()).limit(8).all()
        recent_memories = db.query(Memory).filter(Memory.user_id == user.id).order_by(Memory.id.desc()).limit(5).all()
        recent_reminders = db.query(Reminder).filter(Reminder.user_id == user.id, Reminder.status == "active").order_by(Reminder.id.desc()).limit(5).all()
    context = {
        "first_name": first_name, "branch": branch, "profile_city": city, "profile_state": state,
        "gps_available": lat is not None and lng is not None, "latitude": lat, "longitude": lng, "user_type": user_type,
        "recent_conversation": [{"role": m.role, "content": m.content[:500]} for m in reversed(recent_messages)],
        "memories": [{"title": m.title, "note": (m.note or "")[:300]} for m in recent_memories],
        "reminders": [{"title": r.title, "when": r.when_text} for r in recent_reminders],
    }
    # Deterministic routing comes first so clear requests never fall into a vague AI fallback.
    inferred_intent = infer_intent(message)
    if explicit_intent and explicit_intent != "general":
        plan = {"intent": explicit_intent}
    elif inferred_intent != "general":
        plan = {
            "intent": inferred_intent,
            "needs_places": inferred_intent == "find_local_veteran_activities",
            "needs_location": inferred_intent == "find_local_veteran_activities",
            "search_query": message,
        }
    else:
        plan = await plan_request(message, context)
    intent = plan.get("intent", "general")

    if plan.get("needs_clarification"):
        if any(x in message.lower() for x in ("event", "events", "activity", "activities", "near me")):
            return event_choice_payload(first_name)
        question = clean_text(plan.get("clarification_question")) or f"{first_name}, what should I focus on?"
        default_query = clean_text(plan.get("default_query")) or message
        return {
            "response": question,
            "intent": "clarification",
            "data": {"awaiting_clarification": True, "default_query": default_query, "plan": plan},
        }

    if intent in ("local_search", "find_local_veteran_activities") or plan.get("needs_places"):
        near_me = any(phrase in message.lower() for phrase in ("near me", "nearby", "around me", "close to me"))
        if lat is None and lng is None and (near_me or not clean_text(city)):
            return {
                "response": f"Absolutely {first_name}. What city and state should I search, and are you looking for today, this weekend, or specific dates?",
                "intent": "ask_location",
                "data": {"location_required": True}
            }
        live, items, location_meta = await google_places(city=city, state=state, query=message, lat=lat, lng=lng)
        if location_meta.get("error") == "location_required":
            return {
                "response": f"Absolutely {first_name}. What city and state should I search, and are you looking for today, this weekend, or specific dates?",
                "intent": "ask_location",
                "data": {"location_required": True}
            }
        mode = "live Google Places" if live else "fallback map suggestions"
        place_label = f"near {location_meta.get('city')}, {location_meta.get('state')}" if location_meta.get('city') else "near your current location"
        verified_note = "live options" if live else "search starting points"
        top = items[:3]
        detail_lines = []
        for index, item in enumerate(top, 1):
            title_text = clean_text(item.get("title")) or "Local option"
            location_text = clean_text(item.get("location"))
            rating_text = f", rated {item.get('rating')}" if item.get("rating") else ""
            detail_lines.append(f"{index}. {title_text}" + (f" — {location_text}" if location_text else "") + rating_text)
        listed = " ".join(detail_lines) if detail_lines else "I did not find a verified option yet."
        fallback = (
            f"{first_name}, I found three {verified_note} {place_label}. {listed} "
            "I recommend starting with option 1 because it is the first strong match. "
            "You can choose a number, ask for directions, or say ‘show me free events today.’"
        )
        prompt = f"""User: {first_name}, branch={branch}, current_area={location_meta.get('city') or city}, {location_meta.get('state') or state}
User type: {user_type}
Exact request: {message}
Primary intent: {intent}
Live local results: live={live}, source={location_meta.get('source')}, results={json.dumps(items[:5])}

COMPLETION CONTRACT:
- Do not repeat, quote, paraphrase, or announce the user's request.
- Do not say “I heard you,” “best starting point,” “tell me more,” or “tell me the one detail that matters.”
- Complete the useful work before stopping.
- If live results exist, give the three best options immediately and recommend one.
- If live results are unavailable, clearly say that these are map/search starting points and still give the best three next actions.
- Make the choice easy: number the options and tell the user they can say a number.

Answer ONLY the exact local intent. If the request is for veteran activities or events, discuss only those activities or events. Do not mention documents, reminders, benefits, memories, or other ValorBuddy features.
For each option, briefly state what it is, where it is, and any verified rating or useful live detail available in the results. Do not invent dates, times, hours, admission, or event details that are absent.
Use natural battle-buddy language without calling yourself an AI assistant. Finish with one concrete next action, such as “Say 1, 2, or 3 and I’ll check directions and current details.” Do not ask a vague follow-up question."""
        response = await gemini_reply(prompt, fallback, grounded=False)
        quick_choices = [
            {"label": "Today", "query": "veteran events today near me"},
            {"label": "Free", "query": "free veteran events near me"},
            {"label": "Family-friendly", "query": "family friendly veteran activities near me"},
            {"label": "VFW / Legion", "query": "VFW American Legion events near me"},
            {"label": "Outdoor", "query": "veteran outdoor activities near me"},
            {"label": "Pick for me", "query": "best rated veteran activities closest to me"},
        ]
        return {"response": response, "intent": intent, "data": {"live": live, "items": items, "choices": quick_choices, "location": location_meta, "plan": plan}}

    if intent == "search_benefits":
        data = benefits_lookup(message, state, branch)
        fallback = f"{first_name}, here is the strongest starting point: {data['items'][0]['title']}. {data['items'][0]['summary']} Next step: {data['items'][0]['next_step']}"
        prompt = f"""You are ValorBuddy, a plain-English veteran benefits guide. You are informational only, not legal/medical advice.
User: {first_name}, {branch}, {city}, {state}
Question: {message}
Benefit data: {json.dumps(data)}
Answer naturally and specifically. Include spouse, child, dependent, caregiver, and family access when relevant. Keep it concise and useful."""
        response = await gemini_reply(prompt, fallback, grounded=bool(plan.get("needs_google_search")))
        return {"response": response, "intent": intent, "data": {**data, "plan": plan}}

    if intent == "create_reminder":
        reminder_title = title or message or "Reminder"
        when_text = f"{date} {time}".strip() or "Soon"
        if user and db:
            row = Reminder(user_id=user.id, title=reminder_title, date=date, time=time, when_text=when_text)
            db.add(row); db.commit()
        return {"response": f"Done, {first_name}. I saved this reminder: {reminder_title}. Time: {when_text}.", "intent": intent}

    if intent == "save_memory":
        mem_title = title or "Saved memory"
        note = memory or message
        if user and db:
            db.add(Memory(user_id=user.id, title=mem_title, note=note, tags=["voice", "assistant"])); db.commit()
        return {"response": f"I saved that for you, {first_name}. You can find it in your Memory Wall.", "intent": intent}

    if intent == "suggest_music":
        items = music_suggestions(mood or message or "calm", branch)
        fallback = f"{first_name}, I would start with {items[0]['title']}. If you want, I can also suggest gospel, country, patriotic, or calm focus music."
        prompt = f"User {first_name} asked about music: {message}. Branch: {branch}. Suggestions: {json.dumps(items)}. Respond like a helpful companion with one clear recommendation."
        response = await gemini_reply(prompt, fallback)
        return {"response": response, "intent": intent, "data": {"items": items}}

    if intent == "get_today_briefing":
        rems = []
        if user and db:
            rems = db.query(Reminder).filter(Reminder.user_id == user.id, Reminder.status == "active").order_by(Reminder.id.desc()).limit(3).all()
        live, places, location_meta = await google_places(city=city, state=state, query="veteran friendly coffee parks VFW", lat=lat, lng=lng)
        reminder_txt = "; ".join([f"{r.title} ({r.when_text})" for r in rems]) or "no saved reminders yet"
        fallback = f"Good to see you, {first_name}. You have {reminder_txt}. One nearby option is {places[0]['title']}. Should I check its hours or directions?"
        prompt = f"Create a short daily briefing for {first_name}, a {branch} veteran in {city}, {state}. Reminders: {reminder_txt}. Nearby options: {json.dumps(places[:3])}. Make it warm, practical, and not canned."
        response = await gemini_reply(prompt, fallback)
        return {"response": response, "intent": intent, "data": {"items": places, "live": live}}

    if intent == "get_user_profile":
        return {"response": f"I have your name as {first_name} and your branch as {branch}. Your saved area is {city}, {state}. Tell me what you want corrected, and I’ll focus on that update.", "intent": intent}

    # General companion: make every answer contextual, not canned.
    recent = []
    rems = []
    if user and db:
        recent = db.query(Memory).filter(Memory.user_id == user.id).order_by(Memory.id.desc()).limit(4).all()
        rems = db.query(Reminder).filter(Reminder.user_id == user.id).order_by(Reminder.id.desc()).limit(4).all()
    routed_agents = route_goal(message)
    active_agent = routed_agents[0] if routed_agents else "companion"
    fallback = f"{agent_opening(active_agent, first_name)} I identified one useful next action and kept the response within this agent's mission."
    profile = user.profile if user and user.profile else None
    profile_context = {
        "first_name": first_name, "last_name": getattr(profile, "last_name", ""), "rank": getattr(profile, "rank", ""),
        "branch": branch, "service_status": getattr(profile, "service_status", user_type),
        "service_years": f"{getattr(profile, 'service_start_year', '')}-{getattr(profile, 'service_end_year', '')}",
        "deployment_history": getattr(profile, "deployment_history", ""), "va_rating": getattr(profile, "va_rating", ""),
        "current_city": city, "current_state": state, "interests": getattr(profile, "interests", []),
        "accessibility_needs": getattr(profile, "accessibility_needs", [])
    }
    prompt = build_agent_prompt(active_agent, member=profile_context, request=message, context={"recent_memories": [m.title for m in recent], "recent_reminders": [r.title for r in rems], "primary_intent": intent, "handoff_candidates": routed_agents[1:]}) + """

COMPLETION CONTRACT:
- Never repeat or paraphrase the user's question back to them.
- Never begin with “I heard you,” “you asked,” or “based on what you asked.”
- Do not stop after promising to help. Perform the useful work in the same response.
- Continue until a logical conclusion: answer, recommendation, useful options, and one concrete next action.
- Ask a question only when an essential fact is genuinely missing and cannot be inferred.

Answer the user's actual request directly and stay on that single topic. Do not list ValorBuddy capabilities or unrelated features. Do not say “my next useful step,” “I can help by,” “tell me more,” or “which direction do you want?” Keep it natural, concise, warm, and action-oriented."""
    use_grounding = bool(plan.get("needs_google_search") or intent == "live_web")
    if use_grounding:
        prompt += "\nUse Google Search grounding for current facts. Include dates or freshness context where useful and do not invent local results."
    response = await gemini_reply(prompt, fallback, grounded=use_grounding)
    return {"response": response, "intent": intent, "data": {"plan": plan, "grounded": use_grounding}}


app = FastAPI(title=APP_NAME, version="5.0.0")
origins = [o.strip() for o in CORS_ORIGINS.split(",") if o.strip()]
app.add_middleware(CORSMiddleware, allow_origins=origins or ["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")


@app.on_event("startup")
def startup():
    Base.metadata.create_all(bind=engine)
    # Lightweight additive migration for existing PostgreSQL/SQLite deployments.
    additions = {
        "rank": "VARCHAR(120) DEFAULT ''", "service_status": "VARCHAR(80) DEFAULT 'Veteran'",
        "service_start_year": "VARCHAR(10) DEFAULT ''", "service_end_year": "VARCHAR(10) DEFAULT ''",
        "deployment_history": "TEXT DEFAULT ''", "va_rating": "VARCHAR(30) DEFAULT ''",
        "accessibility_needs": "JSON", "preferred_music_genres": "JSON", "profile_data": "JSON"
    }
    with engine.begin() as conn:
        existing = {c["name"] for c in inspect(engine).get_columns("user_profiles")}
        for name, sql_type in additions.items():
            if name not in existing:
                try:
                    conn.execute(text(f"ALTER TABLE user_profiles ADD COLUMN {name} {sql_type}"))
                except Exception as exc:
                    logger.warning("Profile migration skipped for %s: %s", name, exc)
    db = SessionLocal()
    try:
        admin = db.query(User).filter(User.email == ADMIN_EMAIL).first()
        if not admin and ADMIN_PASSWORD:
            admin = User(email=ADMIN_EMAIL, password_hash=hash_password(ADMIN_PASSWORD), role="admin")
            db.add(admin); db.flush()
            db.add(UserProfile(user_id=admin.id, first_name="Eugene", last_name="Ebem", branch="Army", service_status="Veteran", city="Dallas", state="TX", interests=["administration", "veteran support"]))
        elif admin and admin.role != "admin":
            admin.role = "admin"
        db.commit()
    finally:
        db.close()


@app.get("/health")
def health():
    return {"status": "ok", "app": APP_NAME, "version": "4.9.1", "database": "postgres" if DATABASE_URL.startswith("postgres") else "sqlite", "gemini": bool(GEMINI_API_KEY), "google_places": bool(GOOGLE_MAPS_API_KEY)}


@app.get("/db/tables")
def db_tables():
    return {"tables": sorted(Base.metadata.tables.keys())}


@app.post("/auth/register", response_model=LoginResponse)
def register(payload: RegisterRequest, db: Session = Depends(get_db)):
    if db.query(User).filter(User.email == payload.email).first():
        raise HTTPException(status_code=409, detail="Email already exists")
    user = User(email=str(payload.email).lower(), password_hash=hash_password(payload.password), role="veteran")
    db.add(user); db.flush()
    db.add(UserProfile(user_id=user.id, first_name=payload.first_name, last_name=payload.last_name, rank=payload.rank, branch=payload.branch, service_status=payload.service_status, service_start_year=payload.service_start_year, service_end_year=payload.service_end_year, deployment_history=payload.deployment_history, va_rating=payload.va_rating, city=payload.city, state=payload.state, interests=payload.interests, accessibility_needs=payload.accessibility_needs, preferred_music_genres=payload.preferred_music_genres, profile_data=payload.profile_data))
    db.add(AdminAuditLog(user_id=user.id, action="user.registered", details=user.email))
    db.commit(); db.refresh(user)
    token = create_access_token(user)
    return LoginResponse(token=token, user=profile_out(user))


@app.post("/auth/login", response_model=LoginResponse)
def login(payload: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == str(payload.email).lower()).first()
    if not user or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = create_access_token(user)
    db.add(AuthToken(user_id=user.id, token=token)); db.add(AdminAuditLog(user_id=user.id, action="user.login", details=user.email)); db.commit()
    return LoginResponse(token=token, user=profile_out(user))


@app.get("/auth/me", response_model=ProfileOut)
def me(user: User = Depends(get_current_user)):
    return profile_out(user)


@app.get("/api/profile")
def get_profile(user: User = Depends(get_current_user)):
    return profile_out(user).model_dump()

@app.post("/api/profile/branch")
def update_profile_branch(payload: BranchUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    allowed = {"Army", "Navy", "Air Force", "Marines", "Coast Guard", "Space Force"}
    if payload.branch not in allowed:
        raise HTTPException(status_code=400, detail="Unsupported service branch")
    if not user.profile:
        db.add(UserProfile(user_id=user.id, first_name="Veteran", branch=payload.branch, city="", state=""))
    else:
        user.profile.branch = payload.branch
        user.profile.updated_at = datetime.now(timezone.utc)
    db.add(AdminAuditLog(user_id=user.id, action="profile.branch_updated", details=payload.branch))
    db.commit()
    db.refresh(user)
    return profile_out(user).model_dump()


@app.post("/api/profile")
def update_profile(payload: ProfileUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    p = user.profile or UserProfile(user_id=user.id, first_name=payload.first_name)
    for field in ("first_name", "last_name", "rank", "branch", "service_status", "service_start_year", "service_end_year", "deployment_history", "va_rating", "city", "state", "interests", "accessibility_needs", "preferred_music_genres", "profile_data", "military_mos", "military_job_title", "military_experience", "civilian_career_goal", "business_interest", "military_specialty_description", "years_of_service", "security_clearance", "highest_education", "civilian_certifications", "linkedin_url"):
        setattr(p, field, getattr(payload, field))
    p.updated_at = datetime.now(timezone.utc)
    db.add(p); db.add(AdminAuditLog(user_id=user.id, action="profile.updated", details=f"{payload.first_name} {payload.last_name}".strip()))
    db.commit(); db.refresh(user)
    return profile_out(user).model_dump()


@app.get("/api/events/search")
async def events_search(city: str = "", state: str = "", keyword: str = "veteran events", lat: float | None = None, lng: float | None = None, user: Optional[User] = Depends(get_optional_user), db: Session = Depends(get_db)):
    if user and user.profile:
        city = city or user.profile.city or ""
        state = state or user.profile.state or ""
    live, items, location_meta = await google_places(city=city, state=state, query=keyword, lat=lat, lng=lng)
    db.add(ActivitySearch(user_id=user.id if user else None, city=location_meta.get("city") or city or "", state=location_meta.get("state") or state or "", query=keyword, live=live, results=items)); db.commit()
    return {"live": live, "provider": "Google Places" if live else "Fallback", "items": items, "location": location_meta}


@app.post("/api/reminders")
def create_reminder(payload: ReminderIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    row = Reminder(user_id=user.id, title=payload.title, date=payload.date, time=payload.time, when_text=payload.when_text or f"{payload.date} {payload.time}".strip(), note=payload.note)
    db.add(row); db.add(AdminAuditLog(user_id=user.id, action="reminder.created", details=payload.title)); db.commit(); db.refresh(row)
    return {"id": row.id, "title": row.title, "date": row.date, "time": row.time, "when_text": row.when_text, "status": row.status, "calendar_enabled": GOOGLE_CALENDAR_ENABLED}


@app.get("/api/reminders")
def list_reminders(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(Reminder).filter(Reminder.user_id == user.id).order_by(Reminder.id.desc()).all()
    return [{"id": r.id, "title": r.title, "date": r.date, "time": r.time, "when_text": r.when_text, "status": r.status} for r in rows]


@app.post("/api/memories")
def create_memory(payload: MemoryIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    row = Memory(user_id=user.id, title=payload.title, note=payload.note, tags=payload.tags, image_url=payload.image_url)
    db.add(row); db.add(AdminAuditLog(user_id=user.id, action="memory.created", details=payload.title)); db.commit(); db.refresh(row)
    return {"id": row.id, "title": row.title, "note": row.note, "tags": row.tags, "image_url": row.image_url}


@app.get("/api/memories")
def list_memories(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(Memory).filter(Memory.user_id == user.id).order_by(Memory.id.desc()).all()
    return [{"id": r.id, "title": r.title, "note": r.note, "tags": r.tags, "image_url": r.image_url} for r in rows]


@app.get("/api/benefits/search")
def search_benefits(query: str = "benefits", state: str = "TX", branch: str = "Army"):
    return benefits_lookup(query, state, branch)


@app.get("/api/music/suggest")
def suggest_music(mood: str = "calm", branch: str = "Army"):
    return {"items": music_suggestions(mood, branch)}

@app.get("/api/music/favorites")
def music_favorites(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(MusicFavorite).filter(MusicFavorite.user_id == user.id).order_by(MusicFavorite.id.desc()).all()
    return [{"id": r.id, "title": r.title, "url": r.url, "mood": r.mood} for r in rows]

class MusicFavoriteIn(BaseModel):
    title: str
    url: str = ""
    mood: str = ""

@app.post("/api/music/favorites")
def add_music_favorite(payload: MusicFavoriteIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    row = MusicFavorite(user_id=user.id, title=payload.title, url=payload.url, mood=payload.mood)
    db.add(row); db.commit(); db.refresh(row)
    return {"id": row.id, "title": row.title, "url": row.url, "mood": row.mood}

@app.delete("/api/music/favorites/{favorite_id}")
def delete_music_favorite(favorite_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    row = db.query(MusicFavorite).filter(MusicFavorite.id == favorite_id, MusicFavorite.user_id == user.id).first()
    if not row: raise HTTPException(status_code=404, detail="Music preference not found")
    db.delete(row); db.commit(); return {"deleted": True}


def time_greeting() -> str:
    hour = datetime.now().hour
    return "Good morning" if hour < 12 else ("Good afternoon" if hour < 18 else "Good evening")


@app.get("/api/briefing")
async def today_briefing(lat: float | None = None, lng: float | None = None, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    p = user.profile
    rems = db.query(Reminder).filter(Reminder.user_id == user.id, Reminder.status == "active").order_by(Reminder.id.desc()).limit(3).all()
    live, events, location_meta = await google_places(city=p.city, state=p.state, query="veteran events VA VFW American Legion family friendly", lat=lat, lng=lng)
    return {"greeting": f"{time_greeting()}, {p.first_name}. I’m ValorBuddy. How can I help today?", "location": f"{location_meta.get('city') or p.city}, {location_meta.get('state') or p.state}", "reminders": [{"title": r.title, "when_text": r.when_text} for r in rems], "events": events[:3], "wellness_prompt": ("Live Google Places is connected." if live else "Live place results are unavailable. Check the Google Places configuration.")}


@app.post("/api/companion/chat")
async def companion_chat(payload: CompanionRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    p = user.profile
    conv = db.get(Conversation, payload.conversation_id) if payload.conversation_id else None
    if not conv:
        conv = Conversation(user_id=user.id, source="web", title="ValorBuddy companion")
        db.add(conv); db.flush()
    selected_agents = route_goal(payload.message)
    mission = None
    # Complex, outcome-oriented requests become missions. Casual conversation remains conversational.
    outcome_terms = ("help me", "find", "plan", "build", "create", "prepare", "compare", "apply", "move", "travel", "resume", "business plan", "benefit", "housing", "job", "document")
    should_mission = selected_agents != ["companion"] and (len(payload.message.split()) >= 5 or any(x in payload.message.lower() for x in outcome_terms))
    if should_mission:
        mission = await create_agent_mission(MissionCreateIn(goal=payload.message, lat=payload.lat, lng=payload.lng, timezone=payload.timezone), user, db)
        useful = [step.get("output") for step in mission.get("steps", []) if step.get("output") and step.get("tool_name") not in {"profile.read", "memory.read"}]
        synth_prompt = build_agent_prompt("supervisor", member=profile_out(user).model_dump(), request=payload.message, context={"mission": mission.get("title"), "agents": mission.get("participating_agents"), "verified_results": useful, "next_action": mission.get("next_action")})
        reply = await gemini_reply(synth_prompt, mission.get("summary") or "I created and worked the mission. Open Mission Control to review the completed steps and next action.")
        result = {"mode": "mission", "response": reply, "intent": "agentic_mission", "data": {"mission": mission, "agents": selected_agents}}
    else:
        result = await route_valorbuddy_message(text=payload.message, first_name=p.first_name, branch=p.branch, city=p.city, state=p.state, lat=payload.lat, lng=payload.lng, user=user, db=db)
        reply = result.get("response", "")
    db.add(Message(conversation_id=conv.id, user_id=user.id, role="user", content=payload.message))
    db.add(Message(conversation_id=conv.id, user_id=user.id, role="assistant", content=reply, metadata_json={"intent": result.get("intent"), "data": result.get("data", {})}))
    db.commit()
    return {"conversation_id": conv.id, **result, "reply": reply}



def _extract_document_text(filename: str, content: bytes) -> str:
    # Extract readable text from common veteran documents without OCR hallucination.
    lower = (filename or "").lower()
    try:
        if lower.endswith((".txt", ".md", ".csv", ".json")):
            return content.decode("utf-8", errors="ignore")[:50000]
        if lower.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join((page.extract_text() or "") for page in reader.pages)[:50000]
        if lower.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:50000]
    except Exception as exc:
        logger.warning("Document text extraction failed for %s: %s", filename, exc)
    return ""


def _document_type_from_name(filename: str, requested: str, extracted: str) -> str:
    text = f"{filename} {extracted[:2500]}".lower()
    if any(x in text for x in ["resume", "curriculum vitae", "professional experience", "work experience"]): return "resume"
    if any(x in text for x in ["dd214", "certificate of release", "discharge from active duty"]): return "dd214"
    if any(x in text for x in ["department of veterans affairs", "va decision", "disability rating"]): return "va_record"
    if any(x in text for x in ["certificate", "certification", "license"]): return "certification"
    return requested if requested and requested != "general" else "general"


async def _analyze_document(filename: str, doc_type: str, extracted: str, user: User) -> dict[str, Any]:
    if not extracted.strip():
        return {"classification": doc_type, "summary": "File stored securely. This file did not contain machine-readable text, so OCR or a clearer digital copy is needed for full analysis.", "skills": [], "suggested_actions": ["Open the file to confirm it is readable", "Upload a text-based PDF or DOCX for deeper analysis"]}
    profile = profile_out(user).model_dump()
    shape = {"classification": doc_type, "summary": "", "skills": [], "experience_highlights": [], "education": [], "certifications": [], "suggested_roles": [], "missing_information": [], "suggested_actions": []}
    prompt = f'''Return only valid JSON matching this structure: {json.dumps(shape)}
You are the ValorBuddy Documents Agent and Career Agent working together.
Analyze the uploaded document using only its actual text. Do not invent credentials, dates, employers, MOS codes, clearances, awards, metrics, or education.
Identify transferable skills and useful next actions. If it is a resume, suggest up to five realistic civilian role directions and explain missing information needed to improve it.
Member profile context: {json.dumps(profile, default=str)}
Filename: {filename}
Detected type: {doc_type}
Document text:\n{extracted[:18000]}'''
    raw = await gemini_reply(prompt, json.dumps(shape), json_mode=True)
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else shape
    except Exception:
        return {**shape, "summary": raw[:3000]}

@app.post("/api/documents")
async def upload_document(doc_type: str = Form("general"), file: UploadFile = File(...), user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The selected file is empty")
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File is larger than the 20 MB upload limit")
    original_name = Path(file.filename or "document").name
    safe_name = f"{user.id}_{int(datetime.now().timestamp())}_{re.sub(r'[^A-Za-z0-9._-]+','_',original_name)}"
    path = UPLOAD_DIR / safe_name
    path.write_bytes(content)
    extracted = _extract_document_text(original_name, content)
    detected_type = _document_type_from_name(original_name, doc_type, extracted)
    analysis = await _analyze_document(original_name, detected_type, extracted, user)
    summary = str(analysis.get("summary") or "Document uploaded and indexed.")[:6000]
    row = Document(user_id=user.id, filename=original_name, doc_type=detected_type, file_url=f"/uploads/{safe_name}", extracted_text=extracted, ai_summary=summary, analysis_json=analysis, status="processed" if extracted else "needs_ocr", processed_at=datetime.now(timezone.utc))
    db.add(row); db.flush()
    # Persist high-value extracted facts as reviewable memory, never as invented truth.
    if detected_type == "resume":
        skills = analysis.get("skills") or []
        if skills:
            db.add(MemoryFact(user_id=user.id, category="career", key=f"resume_{row.id}_skills", value=", ".join(map(str, skills[:20])), confidence="extracted", source=f"document:{row.id}"))
    db.add(AdminAuditLog(user_id=user.id, action="document.intelligently_processed", details=f"{detected_type}:{original_name}"))
    db.commit(); db.refresh(row)
    mission = None
    if detected_type in {"resume", "dd214", "certification", "va_record"} and extracted:
        goal = f"Review my uploaded {detected_type} named {original_name}, identify what it means for me, and recommend the strongest next actions."
        try:
            mission = await create_agent_mission(MissionCreateIn(goal=goal, title=f"Analyze {original_name}"), user, db)
        except Exception as exc:
            logger.exception("Automatic document mission failed: %s", exc)
    return {"id": row.id, "filename": row.filename, "doc_type": row.doc_type, "file_url": row.file_url, "ai_summary": row.ai_summary, "analysis": analysis, "status": row.status, "mission": mission}


@app.get("/api/documents")
def list_documents(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(Document).filter(Document.user_id == user.id).order_by(Document.id.desc()).all()
    return [{"id": r.id, "filename": r.filename, "doc_type": r.doc_type, "file_url": r.file_url, "ai_summary": r.ai_summary, "analysis": r.analysis_json or {}, "status": r.status, "processed_at": r.processed_at} for r in rows]


@app.post("/api/vapi/action")
async def vapi_action(payload: VapiActionRequest, db: Session = Depends(get_db)):
    text = payload.message or payload.query or payload.title or payload.memory or ""
    email = (payload.email or "").lower().strip()
    user = db.query(User).filter(User.email == email).first() if email else None
    # Never silently attach an unknown caller to a demo account.
    if not user and email:
        logger.info("Vapi caller not matched to a registered profile: %s", email)

    profile = user.profile if user and user.profile else None
    first_name = payload.first_name or (profile.first_name if profile else "there")
    branch = payload.branch or (profile.branch if profile else "") or "Veteran"

    # IMPORTANT: browser GPS beats profile city. Profile city is used only when GPS is not available.
    city = clean_text(payload.city) or (profile.city if profile else "")
    state = clean_text(payload.state) or (profile.state if profile else "")
    lat = payload.lat
    lng = payload.lng

    result = await route_valorbuddy_message(
        text=text,
        first_name=first_name,
        branch=branch,
        city=city,
        state=state,
        lat=lat,
        lng=lng,
        user_type=payload.user_type or "Veteran",
        user=user,
        db=db,
        explicit_intent=payload.intent,
        title=payload.title,
        date=payload.date,
        time=payload.time,
        memory=payload.memory,
        mood=payload.mood,
        context_items=payload.context_items,
    )
    return result


@app.get("/admin/overview")
def admin_overview(_: User = Depends(admin_required), db: Session = Depends(get_db)):
    return {"users": db.query(User).count(), "veterans": db.query(User).filter(User.role == "veteran").count(), "admins": db.query(User).filter(User.role == "admin").count(), "reminders": db.query(Reminder).count(), "memories": db.query(Memory).count(), "conversations": db.query(Conversation).count(), "messages": db.query(Message).count(), "documents": db.query(Document).count(), "activity_searches": db.query(ActivitySearch).count()}


@app.get("/admin/users")
def admin_users(_: User = Depends(admin_required), db: Session = Depends(get_db)):
    rows = db.query(User).order_by(User.id.desc()).all()
    return [{"id": u.id, "email": u.email, "role": u.role, "active": u.is_active, "first_name": u.profile.first_name if u.profile else "", "last_name": u.profile.last_name if u.profile else "", "rank": u.profile.rank if u.profile else "", "branch": u.profile.branch if u.profile else "", "service_status": u.profile.service_status if u.profile else "", "va_rating": u.profile.va_rating if u.profile else "", "city": u.profile.city if u.profile else "", "state": u.profile.state if u.profile else ""} for u in rows]


@app.get("/admin/activity")
def admin_activity(_: User = Depends(admin_required), db: Session = Depends(get_db)):
    logs = db.query(AdminAuditLog).order_by(AdminAuditLog.id.desc()).limit(100).all()
    return [{"id": l.id, "user_id": l.user_id, "action": l.action, "details": l.details, "created_at": l.created_at.isoformat() if l.created_at else None} for l in logs]


# ===== ValorBuddy v4.9 Platform Edition =====
class MemoryFact(Base):
    __tablename__ = "memory_facts"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    category = Column(String(80), nullable=False, default="preference")
    key = Column(String(120), nullable=False)
    value = Column(Text, nullable=False)
    confidence = Column(String(20), nullable=False, default="confirmed")
    source = Column(String(80), nullable=False, default="member")
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))

class PlatformSetting(Base):
    __tablename__ = "platform_settings"
    id = Column(Integer, primary_key=True)
    setting_key = Column(String(120), unique=True, nullable=False, index=True)
    setting_value = Column(Text, nullable=False, default="")
    category = Column(String(80), nullable=False, default="prompt")
    updated_by = Column(Integer, ForeignKey("users.id"), nullable=True)
    updated_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))

class KnowledgeItem(Base):
    __tablename__ = "knowledge_items"
    id = Column(Integer, primary_key=True)
    category = Column(String(80), nullable=False, index=True)
    title = Column(String(255), nullable=False)
    summary = Column(Text, nullable=False, default="")
    url = Column(String(500), nullable=True)
    tags = Column(JSON, nullable=False, default=list)
    status = Column(String(40), nullable=False, default="published")
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))

class MemoryFactIn(BaseModel):
    category: str = "preference"
    key: str
    value: str
    confidence: str = "confirmed"

class PlatformSettingIn(BaseModel):
    setting_key: str
    setting_value: str
    category: str = "prompt"

class KnowledgeItemIn(BaseModel):
    category: str
    title: str
    summary: str = ""
    url: str = ""
    tags: List[str] = []

ADMIN_ROLES = {"admin", "super_admin", "administrator", "super administrator", "content_manager"}

def platform_admin_required(user: User = Depends(get_current_user)) -> User:
    if user.role.lower() not in ADMIN_ROLES:
        raise HTTPException(status_code=403, detail="Platform administration access required")
    return user

@app.get("/api/platform/capabilities")
def platform_capabilities(user: User = Depends(get_current_user)):
    return {
        "version": "4.9.1",
        "edition": "Agentic Core",
        "role": user.role,
        "can_switch_admin": user.role.lower() in ADMIN_ROLES,
        "modules": ["companion", "travel", "events", "benefits", "va_forms", "housing", "employment", "businesses", "discounts", "financial_education", "vehicles", "music", "memories", "documents", "reminders", "organizations"],
        "integrations": {"gemini": bool(GEMINI_API_KEY), "google_places": bool(GOOGLE_MAPS_API_KEY), "calendar": GOOGLE_CALENDAR_ENABLED, "vapi": bool(os.getenv("VAPI_PUBLIC_KEY")), "firebase": bool(os.getenv("FIREBASE_PROJECT_ID"))}
    }

@app.get("/api/memory/facts")
def list_memory_facts(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(MemoryFact).filter(MemoryFact.user_id == user.id).order_by(MemoryFact.updated_at.desc()).all()
    return [{"id": r.id, "category": r.category, "key": r.key, "value": r.value, "confidence": r.confidence, "source": r.source} for r in rows]

@app.post("/api/memory/facts")
def save_memory_fact(payload: MemoryFactIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    row = db.query(MemoryFact).filter(MemoryFact.user_id == user.id, MemoryFact.key == payload.key).first()
    if not row:
        row = MemoryFact(user_id=user.id, key=payload.key)
    row.category, row.value, row.confidence, row.updated_at = payload.category, payload.value, payload.confidence, datetime.now(timezone.utc)
    db.add(row); db.add(AdminAuditLog(user_id=user.id, action="memory.fact_saved", details=payload.key)); db.commit(); db.refresh(row)
    return {"id": row.id, "category": row.category, "key": row.key, "value": row.value, "confidence": row.confidence}

@app.delete("/api/memory/facts/{fact_id}")
def delete_memory_fact(fact_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    row = db.query(MemoryFact).filter(MemoryFact.id == fact_id, MemoryFact.user_id == user.id).first()
    if not row: raise HTTPException(status_code=404, detail="Memory not found")
    db.delete(row); db.commit(); return {"deleted": True}

@app.get("/api/resources/search")
async def resource_search(category: str, query: str = "", city: str = "", state: str = "", lat: float | None = None, lng: float | None = None, user: User = Depends(get_current_user)):
    profile = user.profile
    city = city or (profile.city if profile else "")
    state = state or (profile.state if profile else "")
    map_queries = {
        "travel": "VA hospital Vet Center veteran friendly hotel fuel rest stop",
        "housing": "veteran housing apartments low credit VA loan housing",
        "employment": "veteran employment center workforce jobs",
        "businesses": "veteran owned business",
        "discounts": "veteran discount",
        "vehicles": "veteran auto dealer military discount",
        "organizations": "VFW American Legion DAV veteran organization",
        "events": "veteran events"
    }
    live, items, location = await google_places(city=city, state=state, query=query or map_queries.get(category, category), lat=lat, lng=lng)
    return {"category": category, "query": query, "live": live, "items": items, "location": location, "disclaimer": "Verify eligibility, availability, pricing, and current terms directly with the provider."}

@app.get("/api/va/forms")
def va_forms(query: str = ""):
    forms = [
        {"form":"VA Form 21-526EZ","title":"Disability Compensation","purpose":"Apply for disability compensation or related benefits.","url":"https://www.va.gov/find-forms/about-form-21-526ez/","documents":["DD214 or separation records","Medical evidence","Dependency records when applicable"]},
        {"form":"VA Form 22-1990","title":"Education Benefits","purpose":"Apply for VA education benefits.","url":"https://www.va.gov/find-forms/about-form-22-1990/","documents":["Service history","School or training program information","Direct deposit information"]},
        {"form":"VA Form 10-10EZ","title":"Health Care Application","purpose":"Apply for VA health care.","url":"https://www.va.gov/find-forms/about-form-10-10ez/","documents":["Military service information","Insurance information","Household financial information when requested"]},
        {"form":"VA Form 26-1880","title":"Certificate of Eligibility","purpose":"Request a VA home loan Certificate of Eligibility.","url":"https://www.va.gov/find-forms/about-form-26-1880/","documents":["Proof of service","Current duty statement when applicable","Discharge documentation"]}
    ]
    q=query.lower().strip()
    if q: forms=[f for f in forms if q in json.dumps(f).lower()] or forms
    return {"items": forms, "submission_status": "not_integrated", "notice": "ValorBuddy can explain and prepare information, but does not claim submission unless an authorized VA integration confirms it."}

@app.get("/api/admin/settings")
def admin_settings(user: User = Depends(platform_admin_required), db: Session = Depends(get_db)):
    rows=db.query(PlatformSetting).order_by(PlatformSetting.category, PlatformSetting.setting_key).all()
    return [{"id":r.id,"setting_key":r.setting_key,"setting_value":r.setting_value,"category":r.category,"updated_at":r.updated_at} for r in rows]

@app.post("/api/admin/settings")
def upsert_admin_setting(payload: PlatformSettingIn, user: User = Depends(platform_admin_required), db: Session = Depends(get_db)):
    row=db.query(PlatformSetting).filter(PlatformSetting.setting_key==payload.setting_key).first() or PlatformSetting(setting_key=payload.setting_key)
    row.setting_value, row.category, row.updated_by, row.updated_at = payload.setting_value, payload.category, user.id, datetime.now(timezone.utc)
    db.add(row); db.add(AdminAuditLog(user_id=user.id, action="platform.setting_updated", details=payload.setting_key)); db.commit(); db.refresh(row)
    return {"id":row.id,"setting_key":row.setting_key,"setting_value":row.setting_value,"category":row.category}

@app.get("/api/admin/knowledge")
def admin_knowledge(user: User = Depends(platform_admin_required), db: Session = Depends(get_db)):
    rows=db.query(KnowledgeItem).order_by(KnowledgeItem.id.desc()).all()
    return [{"id":r.id,"category":r.category,"title":r.title,"summary":r.summary,"url":r.url,"tags":r.tags,"status":r.status} for r in rows]

@app.post("/api/admin/knowledge")
def add_knowledge(payload: KnowledgeItemIn, user: User = Depends(platform_admin_required), db: Session = Depends(get_db)):
    row=KnowledgeItem(category=payload.category,title=payload.title,summary=payload.summary,url=payload.url,tags=payload.tags)
    db.add(row); db.add(AdminAuditLog(user_id=user.id, action="knowledge.created", details=payload.title)); db.commit(); db.refresh(row)
    return {"id":row.id,"category":row.category,"title":row.title,"summary":row.summary,"url":row.url,"tags":row.tags,"status":row.status}

@app.get("/api/admin/analytics")
def admin_analytics(user: User = Depends(platform_admin_required), db: Session = Depends(get_db)):
    intents = db.query(Message.metadata_json, func.count(Message.id)).filter(Message.role == "assistant").group_by(Message.metadata_json).limit(10).all()
    return {
        "members": db.query(User).count(), "conversations": db.query(Conversation).count(), "messages": db.query(Message).count(),
        "voice_sessions": db.query(Conversation).filter(Conversation.source == "voice").count(), "documents": db.query(Document).count(),
        "reminders": db.query(Reminder).count(), "saved_memories": db.query(Memory).count(), "searches": db.query(ActivitySearch).count(),
        "integration_health": {"gemini": bool(GEMINI_API_KEY), "places": bool(GOOGLE_MAPS_API_KEY), "calendar": GOOGLE_CALENDAR_ENABLED}
    }


@app.get("/api/career/documents")
def list_career_documents(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rows = db.query(CareerDocument).filter(CareerDocument.user_id == user.id).order_by(CareerDocument.updated_at.desc()).limit(100).all()
    return [{"id":r.id,"document_type":r.document_type,"title":r.title,"target":r.target,"content":r.content,"status":r.status,"created_at":r.created_at,"updated_at":r.updated_at} for r in rows]


@app.post("/api/career/generate")
async def generate_career_document(payload: CareerGenerateIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    p = user.profile
    if not p:
        raise HTTPException(status_code=400, detail="Complete your service profile first")
    mos = p.military_mos or (p.profile_data or {}).get("mos","")
    military_title = p.military_job_title or (p.profile_data or {}).get("military_job_title","")
    experience = p.military_experience or (p.profile_data or {}).get("military_experience","")
    if not mos and not experience:
        raise HTTPException(status_code=400, detail="Add your MOS/AFSC/Rating and military experience before generating a career document")
    doc_labels={"resume":"Civilian Resume","cover_letter":"Cover Letter","business_plan":"Veteran Business Plan","career_plan":"Career Transition Plan"}
    facts = {
      "name": f"{p.first_name} {p.last_name or ''}".strip(), "branch": p.branch, "rank": p.rank or "",
      "service_status": p.service_status, "service_years": f"{p.service_start_year or ''}-{p.service_end_year or ''}",
      "mos": mos, "military_job_title": military_title, "military_experience": experience,
      "military_specialty_description": p.military_specialty_description or "", "years_of_service": p.years_of_service,
      "civilian_career_goal": p.civilian_career_goal or "", "business_interest": p.business_interest or "",
      "security_clearance": p.security_clearance or "", "highest_education": p.highest_education or "",
      "civilian_certifications": p.civilian_certifications or "", "linkedin_url": p.linkedin_url or "",
      "campaigns": (p.profile_data or {}).get("campaigns",""), "decorations": (p.profile_data or {}).get("decorations",""),
      "target": payload.target, "notes": payload.notes
    }
    prompt = f"""Create a professional {doc_labels[payload.document_type]} for a U.S. military community member.
Use ONLY the verified facts below. Translate military terminology into civilian language, but do not invent employers, dates, degrees, certifications, metrics, clearances, awards, revenue, licenses, or achievements.
For a resume: include summary, transferable skills, military experience, education/certifications placeholders only when not provided, and ATS keywords for the target.
For a business plan: include executive summary, customer problem, solution, veteran/MOS advantage, target market, operations, marketing, revenue model, startup assumptions clearly labeled, risks, milestones, and next steps.
For a cover letter or career plan, use an appropriate professional structure.
Verified profile:
{json.dumps(facts, default=str)}
Return polished plain text with clear headings."""
    fallback = f"""{doc_labels[payload.document_type].upper()}

Target: {payload.target}
Candidate: {facts['name']}
Branch / Rank: {facts['branch']} / {facts['rank']}
MOS / AFSC / Rating: {mos}
Military role: {military_title}

TRANSFERABLE EXPERIENCE
{experience}

NEXT STEP
Review this draft, add verified accomplishments and measurable outcomes, and tailor it to the specific opportunity. No unverified credentials or metrics were added."""
    content = await gemini_reply(prompt, fallback)
    source_doc = db.query(Document).filter(Document.user_id == user.id, Document.doc_type == "resume").order_by(Document.id.desc()).first()
    row = CareerDocument(user_id=user.id, document_type=payload.document_type, title=f"{doc_labels[payload.document_type]} — {payload.target}", target=payload.target, content=content, source_profile=facts, source_document_id=source_doc.id if source_doc else None)
    db.add(row); db.add(AdminAuditLog(user_id=user.id, action="career.document.generated", details=f"{payload.document_type}:{payload.target}"))
    db.commit(); db.refresh(row)
    return {"id":row.id,"document_type":row.document_type,"title":row.title,"target":row.target,"content":row.content,"status":row.status,"created_at":row.created_at}


# ============================================================================
# ValorBuddy v4.9.1 Agentic Core
# Goal -> Plan -> Execute -> Verify -> Remember -> Follow up
# ============================================================================

AGENTIC_CORE_PRINCIPLE = VOS_CORE_PRINCIPLE
AGENT_CATALOG = VOS_AGENT_CATALOG
TOOL_CATALOG = VOS_TOOL_CATALOG


class AgentMission(Base):
    __tablename__ = "agent_missions"
    id = Column(Integer, primary_key=True)
    mission_uid = Column(String(80), unique=True, nullable=False, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    title = Column(String(255), nullable=False)
    goal = Column(Text, nullable=False)
    status = Column(String(40), nullable=False, default="planning", index=True)
    primary_agent = Column(String(80), nullable=False, default="supervisor")
    participating_agents = Column(JSON, nullable=False, default=list)
    plan_json = Column(JSON, nullable=False, default=dict)
    summary = Column(Text, nullable=True)
    next_action = Column(Text, nullable=True)
    progress = Column(Integer, nullable=False, default=0)
    priority = Column(String(30), nullable=False, default="normal")
    risk_level = Column(String(30), nullable=False, default="low")
    context_snapshot = Column(JSON, nullable=False, default=dict)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    updated_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    completed_at = Column(DateTime(timezone=True), nullable=True)


class AgentMissionStep(Base):
    __tablename__ = "agent_mission_steps"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    sequence = Column(Integer, nullable=False)
    agent_key = Column(String(80), nullable=False)
    tool_name = Column(String(120), nullable=False)
    title = Column(String(255), nullable=False)
    status = Column(String(40), nullable=False, default="pending", index=True)
    requires_approval = Column(Boolean, nullable=False, default=False)
    input_json = Column(JSON, nullable=False, default=dict)
    output_json = Column(JSON, nullable=False, default=dict)
    verification_json = Column(JSON, nullable=False, default=dict)
    error_message = Column(Text, nullable=True)
    started_at = Column(DateTime(timezone=True), nullable=True)
    completed_at = Column(DateTime(timezone=True), nullable=True)


class AgentApproval(Base):
    __tablename__ = "agent_approvals"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    step_id = Column(Integer, ForeignKey("agent_mission_steps.id"), nullable=False, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    action_summary = Column(Text, nullable=False)
    status = Column(String(30), nullable=False, default="pending", index=True)
    decision_note = Column(Text, nullable=True)
    requested_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    decided_at = Column(DateTime(timezone=True), nullable=True)


class AgentMissionEvent(Base):
    __tablename__ = "agent_mission_events"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    event_type = Column(String(80), nullable=False)
    message = Column(Text, nullable=False)
    event_data = Column(JSON, nullable=False, default=dict)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class AgentHandoff(Base):
    __tablename__ = "agent_handoffs"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    from_agent = Column(String(80), nullable=False)
    to_agent = Column(String(80), nullable=False)
    reason = Column(Text, nullable=False)
    context_json = Column(JSON, nullable=False, default=dict)
    status = Column(String(40), nullable=False, default="completed")
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class AgentToolRun(Base):
    __tablename__ = "agent_tool_runs"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    step_id = Column(Integer, ForeignKey("agent_mission_steps.id"), nullable=True, index=True)
    agent_key = Column(String(80), nullable=False)
    tool_name = Column(String(120), nullable=False)
    status = Column(String(40), nullable=False, index=True)
    latency_ms = Column(Integer, nullable=True)
    retry_count = Column(Integer, nullable=False, default=0)
    source = Column(String(255), nullable=True)
    error_message = Column(Text, nullable=True)
    input_summary = Column(JSON, nullable=False, default=dict)
    output_summary = Column(JSON, nullable=False, default=dict)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class AgentSource(Base):
    __tablename__ = "agent_sources"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    step_id = Column(Integer, ForeignKey("agent_mission_steps.id"), nullable=True, index=True)
    title = Column(String(500), nullable=False)
    source_url = Column(Text, nullable=True)
    source_type = Column(String(80), nullable=False, default="internal")
    authority_level = Column(String(40), nullable=False, default="unknown")
    verification_status = Column(String(40), nullable=False, default="pending")
    retrieved_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class AgentCheckpoint(Base):
    __tablename__ = "agent_checkpoints"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    step_id = Column(Integer, ForeignKey("agent_mission_steps.id"), nullable=True)
    checkpoint_type = Column(String(80), nullable=False)
    state_json = Column(JSON, nullable=False, default=dict)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class AgentFeedback(Base):
    __tablename__ = "agent_feedback"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    rating = Column(Integer, nullable=True)
    useful = Column(Boolean, nullable=True)
    comment = Column(Text, nullable=True)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class AgentFailure(Base):
    __tablename__ = "agent_failures"
    id = Column(Integer, primary_key=True)
    mission_id = Column(Integer, ForeignKey("agent_missions.id"), nullable=False, index=True)
    step_id = Column(Integer, ForeignKey("agent_mission_steps.id"), nullable=True)
    failure_type = Column(String(80), nullable=False)
    message = Column(Text, nullable=False)
    retryable = Column(Boolean, nullable=False, default=False)
    resolved = Column(Boolean, nullable=False, default=False)
    created_at = Column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    resolved_at = Column(DateTime(timezone=True), nullable=True)


class MissionCreateIn(BaseModel):
    goal: str = Field(min_length=3, max_length=2000)
    title: str = ""
    priority: str = "normal"
    lat: float | None = None
    lng: float | None = None
    timezone: str = ""


class ApprovalDecisionIn(BaseModel):
    approved: bool
    note: str = ""


def _mission_event(db: Session, mission: AgentMission, user_id: int, event_type: str, message: str, data: dict | None = None):
    db.add(AgentMissionEvent(mission_id=mission.id, user_id=user_id, event_type=event_type, message=message, event_data=data or {}))


def _mission_dict(mission: AgentMission, db: Session, include_events: bool = False) -> dict[str, Any]:
    steps = db.query(AgentMissionStep).filter(AgentMissionStep.mission_id == mission.id).order_by(AgentMissionStep.sequence).all()
    approvals = db.query(AgentApproval).filter(AgentApproval.mission_id == mission.id).order_by(AgentApproval.id).all()
    result = {
        "id": mission.id, "mission_uid": mission.mission_uid, "title": mission.title, "goal": mission.goal,
        "status": mission.status, "primary_agent": mission.primary_agent, "participating_agents": mission.participating_agents,
        "summary": mission.summary, "next_action": mission.next_action, "progress": mission.progress,
        "priority": mission.priority, "risk_level": mission.risk_level, "plan": mission.plan_json,
        "created_at": mission.created_at, "updated_at": mission.updated_at, "completed_at": mission.completed_at,
        "steps": [{
            "id": s.id, "sequence": s.sequence, "agent_key": s.agent_key,
            "agent_name": AGENT_CATALOG.get(s.agent_key, {}).get("name", s.agent_key),
            "tool_name": s.tool_name, "title": s.title, "status": s.status,
            "requires_approval": s.requires_approval, "input": s.input_json,
            "output": s.output_json, "verification": s.verification_json, "error": s.error_message,
        } for s in steps],
        "approvals": [{"id": a.id, "step_id": a.step_id, "action_summary": a.action_summary, "status": a.status, "decision_note": a.decision_note} for a in approvals],
    }
    if include_events:
        events = db.query(AgentMissionEvent).filter(AgentMissionEvent.mission_id == mission.id).order_by(AgentMissionEvent.id.desc()).limit(50).all()
        result["events"] = [{"id": e.id, "event_type": e.event_type, "message": e.message, "data": e.event_data, "created_at": e.created_at} for e in events]
    return result


def _fallback_mission_plan(goal: str) -> dict[str, Any]:
    return build_fallback_plan(goal)


async def _plan_agent_mission(goal: str, user: User) -> dict[str, Any]:
    fallback = _fallback_mission_plan(goal)
    p = user.profile
    planner_shape = {
        "title": "Short mission title", "primary_agent": "travel", "participating_agents": ["supervisor", "travel"],
        "risk_level": "low", "success_definition": "Specific measurable completion condition",
        "steps": [{"agent": "supervisor", "tool": "profile.read", "title": "Load context", "input": {}}],
    }
    prompt = f"""Return only valid JSON matching this structure: {json.dumps(planner_shape)}

You are the ValorBuddy Supervisor Agent. Core principle: {AGENTIC_CORE_PRINCIPLE}
Create the smallest useful, safe plan for the member's real-world goal.
Only use agents from: {json.dumps(list(AGENT_CATALOG.keys()))}
Only use tools from: {json.dumps(list(TOOL_CATALOG.keys()))}
Always begin with profile.read and memory.read. Read-only research can execute immediately.
Actions that change data (reminder.create, memory.save) must remain approval-gated.
Never add unrelated agents. Never diagnose, promise VA eligibility, or claim a VA submission.
Member context: branch={getattr(p, 'branch', '')}, service_status={getattr(p, 'service_status', '')}, city={getattr(p, 'city', '')}, state={getattr(p, 'state', '')}, accessibility={getattr(p, 'accessibility_needs', [])}
Goal: {goal}
"""
    raw = await gemini_reply(prompt, json.dumps(fallback), json_mode=True, model=GEMINI_PLANNER_MODEL)
    try:
        plan = json.loads(raw)
        if not isinstance(plan.get("steps"), list) or not plan["steps"]:
            return fallback
        valid_steps = []
        for step in plan["steps"][:18]:
            tool = step.get("tool")
            agent = step.get("agent")
            if tool in TOOL_CATALOG and agent in AGENT_CATALOG:
                valid_steps.append({"agent": agent, "tool": tool, "title": str(step.get("title") or tool)[:255], "input": step.get("input") or {}})
        if not valid_steps:
            return fallback
        plan["steps"] = valid_steps
        plan["participating_agents"] = [a for a in plan.get("participating_agents", []) if a in AGENT_CATALOG] or list(dict.fromkeys(s["agent"] for s in valid_steps))
        plan["primary_agent"] = plan.get("primary_agent") if plan.get("primary_agent") in AGENT_CATALOG else plan["participating_agents"][-1]
        return plan
    except Exception:
        return fallback


async def _execute_agent_tool(tool_name: str, payload: dict[str, Any], mission: AgentMission, user: User, db: Session) -> dict[str, Any]:
    p = user.profile
    if tool_name == "profile.read":
        return {"verified": True, "profile": profile_out(user).model_dump(), "source": "authenticated_profile"}
    if tool_name == "memory.read":
        facts = db.query(MemoryFact).filter(MemoryFact.user_id == user.id).order_by(MemoryFact.updated_at.desc()).limit(25).all()
        return {"verified": True, "facts": [{"category": f.category, "key": f.key, "value": f.value, "confidence": f.confidence} for f in facts], "source": "member_controlled_memory"}
    if tool_name == "resources.search":
        category = payload.get("category", "events")
        live, items, location = await google_places(city=getattr(p, "city", ""), state=getattr(p, "state", ""), query=payload.get("query") or category, lat=mission.context_snapshot.get("lat"), lng=mission.context_snapshot.get("lng"))
        return {"verified": bool(items), "live": live, "items": items[:6], "location": location, "source": "google_places" if live else "fallback_search"}
    if tool_name == "benefits.guide":
        data = benefits_lookup(payload.get("query", mission.goal), getattr(p, "state", ""), getattr(p, "branch", ""))
        return {"verified": True, **data, "source": "valorbuddy_benefits_guide", "notice": "Educational guidance; verify eligibility with VA or an accredited representative."}
    if tool_name == "va_forms.search":
        q = (payload.get("query") or mission.goal).lower()
        all_forms = va_forms(q)
        return {"verified": True, **all_forms, "source": "official_va_links"}
    if tool_name in {"travel.search", "housing.search", "employment.search", "discounts.search", "vehicle.research"}:
        category = tool_name.split(".")[0]
        query = payload.get("query") or mission.goal
        search_query = {"travel": f"{query} VA hospitals veteran friendly hotels rest stops fuel safety", "housing": f"{query} veteran friendly housing accessible apartments VA loan resources", "employment": f"{query} veteran hiring jobs apprenticeships federal remote", "discounts": f"{query} verified veteran military discounts", "vehicle": f"{query} veteran vehicle incentives financing insurance EV incentives"}.get(category, query)
        live, items, location = await google_places(city=getattr(p, "city", ""), state=getattr(p, "state", ""), query=search_query, lat=mission.context_snapshot.get("lat"), lng=mission.context_snapshot.get("lng"))
        return {"verified": bool(items), "category": category, "live": live, "items": items[:8], "location": location, "source": "google_places" if live else "search_starting_points"}
    if tool_name == "career.generate":
        p = user.profile
        mos = (p.military_mos or (p.profile_data or {}).get("mos","")) if p else ""
        exp = (p.military_experience or (p.profile_data or {}).get("military_experience","")) if p else ""
        target = payload.get("target") or payload.get("query") or mission.goal
        doc_type = "business_plan" if "business" in target.lower() else "resume"
        req = CareerGenerateIn(document_type=doc_type, target=target[:255], notes="Generated from Agentic Mission Control")
        result = await generate_career_document(req, user, db)
        return {"verified": True, "document": result, "mos": mos, "source": "mos_career_agent"}

    if tool_name == "documents.review":
        rows = db.query(Document).filter(Document.user_id == user.id).order_by(Document.id.desc()).limit(20).all()
        return {"verified": True, "items": [{"id": r.id, "title": r.filename, "status": r.status, "doc_type": r.doc_type} for r in rows], "source": "secure_document_metadata"}
    if tool_name == "finance.educate":
        q=(payload.get("query") or mission.goal).lower()
        topic="budgeting" if "budget" in q else "credit" if "credit" in q else "retirement" if "retire" in q else "investing basics" if "invest" in q else "financial readiness"
        return {"verified": True, "topic": topic, "guidance": ["Clarify the goal and time horizon", "Review cash flow, obligations, and emergency reserves", "Compare options, fees, risks, and official resources", "Choose one small next action"], "notice": "General education only; not individualized financial, tax, legal, or investment advice.", "source": "valorbuddy_financial_education"}
    if tool_name == "family.plan":
        return {"verified": True, "plan": ["Identify the person or commitment that matters", "Choose a specific connection action", "Set a reminder only with member approval"], "source": "valorbuddy_family_support"}
    if tool_name == "wellness.support":
        prompt = build_agent_prompt("wellness", member=profile_out(user).model_dump(), request=payload.get("query") or mission.goal, context={"mode":"non-clinical wellness"})
        response = await gemini_reply(prompt, "Pause, take one slow breath, reduce the next task to one manageable step, and contact a trusted person or professional when more support is needed.")
        return {"verified": True, "response": response, "source": "valorbuddy_wellness"}
    if tool_name == "entertainment.suggest":
        items = music_suggestions(payload.get("query") or mission.goal, getattr(p, "branch", "Army"))
        return {"verified": True, "items": items, "source": "personalized_suggestions"}
    if tool_name == "companion.support":
        mode = payload.get("mode", "companion")
        agent = "supervisor" if mode == "synthesis" else ("safety" if mode == "safety" else "companion")
        completed = [row.output_json for row in db.query(AgentMissionStep).filter(AgentMissionStep.mission_id == mission.id, AgentMissionStep.status == "completed").all()]
        prompt = build_agent_prompt(agent, member=profile_out(user).model_dump(), request=payload.get("query") or mission.goal, context={"mode": mode, "completed_steps": completed})
        response = await gemini_reply(prompt, "Choose the single next action that reduces the most stress or saves the most time, and complete that before expanding the mission.")
        return {"verified": True, "response": response, "agent": agent, "source": f"valorbuddy_{agent}"}
    if tool_name == "reminder.create":
        row = Reminder(user_id=user.id, title=payload.get("title") or mission.title, when_text=payload.get("when_text") or "To be confirmed", note=f"Created from mission {mission.mission_uid}")
        db.add(row); db.flush()
        return {"verified": True, "reminder_id": row.id, "title": row.title, "when_text": row.when_text, "source": "valorbuddy_reminders"}
    if tool_name == "memory.save":
        row = MemoryFact(user_id=user.id, category=payload.get("category", "mission"), key=payload.get("key", f"mission_{mission.id}"), value=payload.get("value", mission.goal), confidence="confirmed", source="approved_agent_mission")
        db.add(row); db.flush()
        return {"verified": True, "memory_fact_id": row.id, "key": row.key, "value": row.value, "source": "member_approved_memory"}
    raise ValueError(f"Unsupported tool: {tool_name}")


async def _run_mission(mission: AgentMission, user: User, db: Session) -> AgentMission:
    steps = db.query(AgentMissionStep).filter(AgentMissionStep.mission_id == mission.id).order_by(AgentMissionStep.sequence).all()
    mission.status = "executing"
    _mission_event(db, mission, user.id, "execution_started", "ValorBuddy started working the mission.")
    db.commit()

    for step in steps:
        if step.status in {"completed", "skipped"}:
            continue
        if step.requires_approval:
            approval = db.query(AgentApproval).filter(AgentApproval.step_id == step.id).first()
            if not approval:
                approval = AgentApproval(mission_id=mission.id, step_id=step.id, user_id=user.id, action_summary=step.title)
                db.add(approval)
            if approval.status == "denied":
                step.status = "skipped"
                _mission_event(db, mission, user.id, "step_skipped", f"Skipped: {step.title}")
                db.commit()
                continue
            if approval.status != "approved":
                step.status = "waiting_for_approval"
                mission.status = "waiting_for_approval"
                mission.next_action = f"Review and approve: {step.title}"
                _mission_event(db, mission, user.id, "approval_requested", mission.next_action, {"step_id": step.id})
                db.commit()
                break
        step.status = "running"; step.started_at = datetime.now(timezone.utc)
        tool_run = AgentToolRun(mission_id=mission.id, step_id=step.id, agent_key=step.agent_key, tool_name=step.tool_name, status="running", input_summary={"keys": list((step.input_json or {}).keys())})
        db.add(tool_run); db.commit()
        started_perf = time.perf_counter()
        try:
            output = await _execute_agent_tool(step.tool_name, step.input_json or {}, mission, user, db)
            tool_run.status = "completed"
            tool_run.latency_ms = int((time.perf_counter() - started_perf) * 1000)
            tool_run.source = output.get("source", "internal")
            tool_run.output_summary = {"verified": bool(output.get("verified", False)), "keys": list(output.keys())[:20]}
            step.output_json = output
            verified = bool(output.get("verified", False))
            step.verification_json = {"verified": verified, "checked_at": datetime.now(timezone.utc).isoformat(), "source": output.get("source", "internal")}
            step.status = "completed" if verified else "failed"
            step.completed_at = datetime.now(timezone.utc)
            _mission_event(db, mission, user.id, "step_completed" if verified else "step_failed", f"{step.title}: {'completed' if verified else 'could not be verified'}", {"step_id": step.id, "tool": step.tool_name})
            if not verified:
                mission.status = "needs_attention"; mission.next_action = f"Review the result for: {step.title}"; db.commit(); break
            db.commit()
        except Exception as exc:
            logger.exception("Agent tool failed: %s", exc)
            tool_run.status = "failed"; tool_run.latency_ms = int((time.perf_counter() - started_perf) * 1000); tool_run.error_message = str(exc)[:1000]
            db.add(AgentFailure(mission_id=mission.id, step_id=step.id, failure_type="tool_execution", message=str(exc)[:2000], retryable=True))
            step.status = "failed"; step.error_message = str(exc)[:1000]; step.completed_at = datetime.now(timezone.utc)
            mission.status = "needs_attention"; mission.next_action = f"Retry or revise: {step.title}"
            _mission_event(db, mission, user.id, "step_failed", f"The step could not be completed: {step.title}", {"error": str(exc)[:500]})
            db.commit(); break

    steps = db.query(AgentMissionStep).filter(AgentMissionStep.mission_id == mission.id).order_by(AgentMissionStep.sequence).all()
    completed = sum(1 for s in steps if s.status in {"completed", "skipped"})
    mission.progress = int((completed / max(len(steps), 1)) * 100)
    if all(s.status in {"completed", "skipped"} for s in steps):
        mission.status = "completed"; mission.progress = 100; mission.completed_at = datetime.now(timezone.utc)
        mission.next_action = "Mission complete. Review the results and start the next mission when ready."
        useful = [s.output_json for s in steps if s.output_json and s.tool_name not in {"profile.read", "memory.read"}]
        mission.summary = "Mission completed with verified steps." if useful else "Mission context reviewed and organized."
        _mission_event(db, mission, user.id, "mission_completed", "Mission completed and verified.")
    mission.updated_at = datetime.now(timezone.utc)
    db.commit(); db.refresh(mission)
    return mission


@app.get("/api/agentic/core")
def agentic_core_info(user: User = Depends(get_current_user)):
    return {
        "version": "5.0.0", "name": "ValorBuddy Veteran Operating System", "core_principle": AGENTIC_CORE_PRINCIPLE,
        "operating_model": "Goal → Plan → Execute → Verify → Remember → Follow up",
        "agents": AGENT_CATALOG, "tools": TOOL_CATALOG,
        "guardrails": ["Authenticated member context", "Approval before consequential actions", "Verified tool results", "No false VA submissions", "User-controlled memory", "Audit trail"],
    }


@app.post("/api/agentic/missions")
async def create_agent_mission(payload: MissionCreateIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    plan = await _plan_agent_mission(payload.goal, user)
    mission = AgentMission(
        mission_uid=f"VB-{datetime.now(timezone.utc).strftime('%Y%m%d')}-{secrets.token_hex(4).upper()}",
        user_id=user.id, title=(payload.title or plan.get("title") or payload.goal[:70])[:255], goal=payload.goal,
        status="planned", primary_agent=plan.get("primary_agent", "supervisor"),
        participating_agents=plan.get("participating_agents", ["supervisor"]), plan_json=plan,
        priority=payload.priority, risk_level=plan.get("risk_level", "low"),
        context_snapshot={"lat": payload.lat, "lng": payload.lng, "timezone": payload.timezone, "created_from": "mission_control"},
        next_action="ValorBuddy is ready to execute the plan.",
    )
    db.add(mission); db.flush()
    for i, item in enumerate(plan.get("steps", []), 1):
        tool = item["tool"]
        db.add(AgentMissionStep(
            mission_id=mission.id, sequence=i, agent_key=item["agent"], tool_name=tool,
            title=item.get("title", tool), requires_approval=bool(TOOL_CATALOG.get(tool, {}).get("approval")), input_json=item.get("input") or {},
        ))
    _mission_event(db, mission, user.id, "mission_created", "Mission planned by the Supervisor Agent.", {"agents": mission.participating_agents})
    db.commit(); db.refresh(mission)
    mission = await _run_mission(mission, user, db)
    return _mission_dict(mission, db, include_events=True)


@app.get("/api/agentic/missions")
def list_agent_missions(status: str = "", user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    query = db.query(AgentMission).filter(AgentMission.user_id == user.id)
    if status:
        query = query.filter(AgentMission.status == status)
    rows = query.order_by(AgentMission.updated_at.desc()).limit(100).all()
    return [_mission_dict(row, db) for row in rows]


@app.get("/api/agentic/missions/{mission_id}")
def get_agent_mission(mission_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    mission = db.query(AgentMission).filter(AgentMission.id == mission_id, AgentMission.user_id == user.id).first()
    if not mission: raise HTTPException(status_code=404, detail="Mission not found")
    return _mission_dict(mission, db, include_events=True)


@app.post("/api/agentic/missions/{mission_id}/run")
async def run_agent_mission(mission_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    mission = db.query(AgentMission).filter(AgentMission.id == mission_id, AgentMission.user_id == user.id).first()
    if not mission: raise HTTPException(status_code=404, detail="Mission not found")
    if mission.status in {"completed", "cancelled"}: return _mission_dict(mission, db, include_events=True)
    mission = await _run_mission(mission, user, db)
    return _mission_dict(mission, db, include_events=True)


@app.post("/api/agentic/approvals/{approval_id}")
async def decide_agent_approval(approval_id: int, payload: ApprovalDecisionIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    approval = db.query(AgentApproval).filter(AgentApproval.id == approval_id, AgentApproval.user_id == user.id).first()
    if not approval: raise HTTPException(status_code=404, detail="Approval request not found")
    if approval.status != "pending": raise HTTPException(status_code=409, detail="Approval has already been decided")
    approval.status = "approved" if payload.approved else "denied"; approval.decision_note = payload.note; approval.decided_at = datetime.now(timezone.utc)
    step = db.query(AgentMissionStep).filter(AgentMissionStep.id == approval.step_id).first()
    mission = db.query(AgentMission).filter(AgentMission.id == approval.mission_id).first()
    if not payload.approved and step: step.status = "skipped"
    _mission_event(db, mission, user.id, "approval_decided", f"Action {'approved' if payload.approved else 'declined'}: {approval.action_summary}")
    db.commit()
    mission = await _run_mission(mission, user, db)
    return _mission_dict(mission, db, include_events=True)


@app.post("/api/agentic/missions/{mission_id}/feedback")
def submit_agent_feedback(mission_id: int, payload: MissionFeedbackIn, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    mission = db.query(AgentMission).filter(AgentMission.id == mission_id, AgentMission.user_id == user.id).first()
    if not mission:
        raise HTTPException(status_code=404, detail="Mission not found")
    row = AgentFeedback(mission_id=mission.id, user_id=user.id, rating=payload.rating, useful=payload.useful, comment=payload.comment)
    db.add(row)
    _mission_event(db, mission, user.id, "feedback_received", "Member feedback recorded.", {"rating": payload.rating, "useful": payload.useful})
    db.commit()
    return {"ok": True, "feedback_id": row.id}


@app.post("/api/agentic/missions/{mission_id}/cancel")
def cancel_agent_mission(mission_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    mission = db.query(AgentMission).filter(AgentMission.id == mission_id, AgentMission.user_id == user.id).first()
    if not mission: raise HTTPException(status_code=404, detail="Mission not found")
    mission.status = "cancelled"; mission.next_action = "Mission cancelled by the member."; mission.updated_at = datetime.now(timezone.utc)
    _mission_event(db, mission, user.id, "mission_cancelled", "Mission cancelled by the member.")
    db.commit(); return _mission_dict(mission, db, include_events=True)
